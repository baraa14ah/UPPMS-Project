#!/usr/bin/env python3
"""Generate Chapter 7 (Testing & Evaluation) — full English edition."""
from __future__ import annotations

from pathlib import Path

from playwright.sync_api import sync_playwright

ROOT = Path(__file__).resolve().parents[1] / "chapter-07-testing-and-evaluation"
OUT = ROOT / "sections-en"
DIAG = ROOT / "diagrams-en"
OUT.mkdir(parents=True, exist_ok=True)
DIAG.mkdir(parents=True, exist_ok=True)

TOTAL_TESTS = 141
UNIT_TESTS = 50
FEATURE_TESTS = 91

STYLE = """
  @page { size: A4; margin: 20mm 16mm 18mm 16mm; }
  body {
    font-family: "Times New Roman", Times, serif;
    font-size: 12.5pt; line-height: 1.7; color: #111; margin: 0; background: #fff;
    text-align: justify; direction: ltr;
  }
  .head {
    border-bottom: 2px solid #111; padding-bottom: 10px; margin-bottom: 22px;
  }
  .head .chapter { font-size: 11pt; color: #444; margin: 0 0 6px; }
  .head h1 { font-size: 17pt; margin: 0; font-weight: bold; text-align: left; }
  h2 { font-size: 13.5pt; margin: 18px 0 10px; page-break-after: avoid; }
  h3 { font-size: 12.5pt; margin: 14px 0 8px; page-break-after: avoid; }
  p { margin: 0 0 11px; }
  ul, ol { margin: 6px 0 14px; padding-left: 1.4em; }
  li { margin: 3px 0; }
  table {
    width: 100%; border-collapse: collapse; margin: 12px 0 16px; font-size: 10.5pt;
    page-break-inside: avoid;
  }
  th, td { border: 1px solid #222; padding: 7px 8px; text-align: left; vertical-align: top; }
  th { background: #f3f3f3; font-weight: bold; }
  .imgbox { margin: 16px 0; text-align: center; page-break-inside: avoid; }
  .imgbox img { max-width: 100%; height: auto; border: 1px solid #ccc; }
  .caption { font-size: 11pt; text-align: center; margin: 8px 0 16px; color: #333; }
  .note {
    margin: 12px 0; padding: 8px 12px; border: 1px solid #777; font-size: 11.5pt;
  }
  .file-ref { font-size: 11pt; color: #444; margin: 8px 0 16px; }
  code { font-family: Consolas, "Courier New", monospace; font-size: 10.5pt; }
"""


def wrap(section_no: str, title: str, body: str) -> str:
    return f"""<!DOCTYPE html>
<html lang="en" dir="ltr">
<head>
  <meta charset="UTF-8" />
  <title>{section_no} — {title}</title>
  <style>{STYLE}</style>
</head>
<body>
  <div class="head">
    <p class="chapter">Chapter Seven — Testing and Evaluation · UPPMS</p>
    <h1>{section_no} {title}</h1>
  </div>
  {body}
</body>
</html>
"""


def img(name: str, caption: str) -> str:
    path = DIAG / name
    if path.exists():
        return f"""
<div class="imgbox"><img src="{path.as_uri()}" alt="{caption}" /></div>
<p class="caption">{caption}</p>
<p class="file-ref">Image file: <code>diagrams-en/{name}</code></p>
"""
    return f'<p class="note">Diagram: <code>diagrams-en/{name}</code></p>'


SECTIONS: list[tuple[str, str, str, str]] = [
    (
        "01",
        "7-1",
        "Test Plan",
        f"""
<p>
This chapter aims to demonstrate that UPPMS works according to its functional and non-functional requirements.
The test plan defines <em>what</em> is tested, <em>when</em>, with which tools, and who is responsible.
</p>

<h2>1) Test Scope</h2>
<ul>
  <li>Laravel API and Services layer (backend).</li>
  <li>University isolation rules (Multi-Tenancy / TenantScope).</li>
  <li>XML registration, proposals, academic tracks, committees, and genetic scheduling.</li>
  <li>Acceptance scenarios on Syrian Private University campus demo data (SPU Demo).</li>
</ul>

<h2>2) Test Environment</h2>
<table>
  <tr><th>Item</th><th>Details</th></tr>
  <tr><td>Test framework</td><td>PHPUnit 10 (Laravel)</td></tr>
  <tr><td>Command</td><td><code>php vendor/bin/phpunit</code> from <code>backend last/</code></td></tr>
  <tr><td>Database</td><td>MySQL (test connection via Laravel / phpunit.xml)</td></tr>
  <tr><td>Frontend</td><td>Manual / acceptance testing on React SPA (Vite)</td></tr>
</table>

<h2>3) Schedule and Responsibilities</h2>
<table>
  <tr><th>Phase</th><th>Owner</th><th>Activity</th></tr>
  <tr><td>During each Sprint</td><td>Development team</td><td>Unit + Feature tests for the related feature</td></tr>
  <tr><td>End of phase</td><td>Development team</td><td>Full suite run + Regression</td></tr>
  <tr><td>Before delivery</td><td>Team + pilot users</td><td>Acceptance / User Testing on SPU Demo</td></tr>
</table>

<h2>4) Plan Outputs</h2>
<ul>
  <li>Documented test cases under <code>tests/Unit</code> and <code>tests/Feature</code>.</li>
  <li>Execution reports (pass / fail / duration).</li>
  <li>Quality metrics and user feedback.</li>
</ul>

<div class="note">
Detailed test-case listings are completed in <strong>Appendix C</strong> for the final report submission.
</div>

{img("01-test-plan-flow.png", "Figure 7-1: Test Plan Flow")}
""",
    ),
    (
        "02",
        "7-2",
        "Types of Tests",
        f"""
<p>
The project applies the testing types required by the graduation-project report guide,
with automated backend focus and manual acceptance on the UI.
</p>

<table>
  <tr><th>Test Type</th><th>Application in UPPMS</th><th>Tool / Approach</th></tr>
  <tr><td>Unit Testing</td><td>FitnessCalculator, EvolutionaryOperators, TrackService, XmlImportService, AI services</td><td>PHPUnit Unit ({UNIT_TESTS} cases)</td></tr>
  <tr><td>Integration Testing</td><td>Services with database and Eloquent models</td><td>Feature/Unit with DatabaseTransactions</td></tr>
  <tr><td>System Testing</td><td>Full API flows: proposals, tracks, committees, XML, scheduling</td><td>PHPUnit Feature ({FEATURE_TESTS} cases)</td></tr>
  <tr><td>Acceptance Testing</td><td>Admin / Supervisor / Student scenarios on SPU campus</td><td>Manual runs + SpuCampusDemoSeeder data</td></tr>
  <tr><td>Performance / Stress</td><td>Genetic scheduling engine on dozens of projects</td><td>GA runs on Campus Demo dataset</td></tr>
  <tr><td>Security Testing</td><td>Tenant isolation, RBAC, XML matching, pending-user blocking</td><td>Feature tests + Middleware review</td></tr>
  <tr><td>Usability Testing</td><td>Clarity of proposals, track timeline, scheduling, XML import UI</td><td>User testing + notes</td></tr>
  <tr><td>Compatibility Testing</td><td>Modern browsers (Chrome/Edge) and RTL/LTR modes</td><td>Manual UI verification</td></tr>
  <tr><td>Regression Testing</td><td>Re-run suite after major fixes</td><td><code>php vendor/bin/phpunit</code></td></tr>
</table>

<p>Total automated test cases in the repository: <strong>{TOTAL_TESTS}</strong> (Unit: {UNIT_TESTS} — Feature: {FEATURE_TESTS}).</p>

{img("02-test-pyramid.png", "Figure 7-2: UPPMS Test Pyramid")}
""",
    ),
    (
        "03",
        "7-3",
        "Test Diagrams",
        f"""
<p>
Test diagrams illustrate the execution plan, module coverage, and result distribution.
</p>

<h2>1) Module Coverage Diagram</h2>
{img("03-test-coverage-modules.png", "Figure 7-3: Modules Covered by Test Cases")}

<h2>2) Sample Primary Test Cases</h2>
<table>
  <tr><th>ID</th><th>Module</th><th>Objective</th><th>Expected Result</th></tr>
  <tr><td>TC-XML-01</td><td>XmlRegistration</td><td>Student registers with matching email + university number from XML</td><td>Account becomes active</td></tr>
  <tr><td>TC-XML-02</td><td>XmlRegistration</td><td>Supervisor registers with matching email only</td><td>Account becomes active</td></tr>
  <tr><td>TC-PROP-01</td><td>ProjectProposal</td><td>Submit up to 3 proposals without early track lock</td><td>Submission succeeds without premature track assignment</td></tr>
  <tr><td>TC-PROP-02</td><td>ProjectProposal</td><td>Approve a proposal</td><td>Project is created + track is assigned</td></tr>
  <tr><td>TC-TRACK-01</td><td>Track</td><td>Prevent skipping prerequisites</td><td>Stage remains locked / request rejected</td></tr>
  <tr><td>TC-GA-01</td><td>GeneticScheduler</td><td>Generate schedules without hard conflicts</td><td>Valid candidates with fitness scores</td></tr>
  <tr><td>TC-COM-01</td><td>Committee</td><td>Prevent project supervisor from sitting on that project's committee</td><td>Conflict rejected</td></tr>
</table>

<h2>3) Results Diagram</h2>
{img("04-results-summary.png", "Figure 7-4: Automated Test Case Distribution")}

<div class="note">
Execution snapshots can be attached in Appendix C with PHPUnit output.
</div>
""",
    ),
    (
        "04",
        "7-4",
        "Test Results and Analysis",
        f"""
<p>
Analysis is based on the automated test inventory in the repository, PHPUnit execution,
and acceptance scenarios on SPU demo data.
</p>

<h2>1) Automated Inventory Summary</h2>
<table>
  <tr><th>Metric</th><th>Value</th></tr>
  <tr><td>Total test cases (list-tests)</td><td>{TOTAL_TESTS}</td></tr>
  <tr><td>Unit</td><td>{UNIT_TESTS}</td></tr>
  <tr><td>Feature / System</td><td>{FEATURE_TESTS}</td></tr>
  <tr><td>Test files</td><td>20 files (Unit + Feature)</td></tr>
</table>

<h2>2) Coverage by Main Test Files</h2>
<table>
  <tr><th>File</th><th>Type</th><th>Verification Focus</th></tr>
  <tr><td>ProjectProposalControllerTest</td><td>Feature</td><td>Submit / approve / delete proposals and tenant isolation</td></tr>
  <tr><td>TrackControllerTest / TrackServiceTest</td><td>Feature + Unit</td><td>Tracks, progress, and prerequisites</td></tr>
  <tr><td>XmlImport* / XmlRegistrationTest</td><td>Feature + Unit</td><td>Import and registration matching</td></tr>
  <tr><td>CommitteeControllerTest</td><td>Feature</td><td>Committee management and conflicts</td></tr>
  <tr><td>GeneticScheduler* / Fitness* / Evolutionary*</td><td>Unit</td><td>Scheduling engine and constraints</td></tr>
  <tr><td>ScheduleControllerTest / ScheduleApproval*</td><td>Feature + Unit</td><td>Generation and approval</td></tr>
  <tr><td>AIIdeation* / AITask*</td><td>Unit + Feature</td><td>AI services</td></tr>
</table>

<h2>3) Results Analysis</h2>
<ul>
  <li><strong>Passed cases:</strong> cover critical paths (XML, proposals, tracks, committees, scheduling).</li>
  <li><strong>Failures / errors:</strong> fixed immediately then re-validated with Regression; logged when they appear in full runs.</li>
  <li><strong>Remaining defects:</strong> no critical security or tenant-isolation defects are accepted before delivery.</li>
  <li><strong>Execution time:</strong> depends on machine and test DB; recorded from PHPUnit output for formal runs.</li>
</ul>

<div class="note">
Official documentation command:
<code>cd "backend last" &amp;&amp; php vendor/bin/phpunit --testdox</code>
</div>
""",
    ),
    (
        "05",
        "7-5",
        "Quality Metrics",
        f"""
<p>
The following metrics are used to evaluate UPPMS quality after testing.
</p>

<table>
  <tr><th>Metric</th><th>Definition</th><th>Measurement in the Project</th></tr>
  <tr><td>Test Coverage (functional)</td><td>Share of critical modules/features covered by tests</td><td>Core modules covered: XML, Proposals, Tracks, Committees, GA, AI</td></tr>
  <tr><td>Defect Density</td><td>Defects found ÷ component size</td><td>Defects logged per Spec Kit phase and re-tested</td></tr>
  <tr><td>Pass Rate</td><td>Passed cases ÷ total executed cases</td><td>Delivery target: ≥ 95% on critical paths</td></tr>
  <tr><td>Response Time</td><td>API latency for common operations</td><td>CRUD/proposals feel interactive in UI; GA duration grows with input size</td></tr>
  <tr><td>Error Rate</td><td>Share of incorrect responses / unhandled exceptions</td><td>Monitored via failing tests and Laravel logs</td></tr>
  <tr><td>Security Isolation</td><td>No cross-university data leakage</td><td>Verified via TenantScope and multi-university Feature tests</td></tr>
</table>

<h2>Quality Acceptance Criteria</h2>
<ul>
  <li>No failures in security, isolation, or XML registration tests.</li>
  <li>Successful basic acceptance scenarios for Admin / Supervisor / Student.</li>
  <li>Ability to generate and approve a defense schedule on Campus Demo without visible hard conflicts.</li>
</ul>
""",
    ),
    (
        "06",
        "7-6",
        "User Testing",
        f"""
<p>
User testing is performed on a Syrian Private University pilot environment using
<code>SpuCampusDemoSeeder</code> data (28 students, 10 supervisors, 28 projects, 4 committees).
</p>

<h2>1) User Sample and Roles</h2>
<table>
  <tr><th>Role</th><th>Demo Account</th><th>Test Scenario</th></tr>
  <tr><td>University Admin</td><td>spu-campus-admin@syrian-private.local</td><td>XML import, tracks, committees, schedule generate &amp; approve</td></tr>
  <tr><td>Supervisor</td><td>spu-campus-supervisor-01@syrian-private.local</td><td>Review proposals, availability, follow projects</td></tr>
  <tr><td>Student</td><td>spu-campus-student-01@syrian-private.local</td><td>Submit proposal, track progress, tasks, notifications</td></tr>
</table>
<p>Shared demo password: <code>password</code></p>

<h2>2) Survey / Interview Themes</h2>
<ul>
  <li>Ease of registration and match with official university records.</li>
  <li>Clarity of the proposal flow until project creation.</li>
  <li>Clarity of the academic track (locked / open stages).</li>
  <li>Usefulness of AI ideation and task generation.</li>
  <li>Clarity of scheduling results for committees and admin.</li>
  <li>Suggested improvements.</li>
</ul>

<h2>3) Initial Findings from Internal Pilots</h2>
<table>
  <tr><th>Observation</th><th>Impact</th><th>Improvement</th></tr>
  <tr><td>Early track assignment on proposal submit caused premature locking</td><td>Confusion in stage UI</td><td>Defer track assignment until approval (fixed)</td></tr>
  <tr><td>Multi-university supervisors might miss proposals for membership universities</td><td>Delayed review</td><td>Extend TenantScope to supervisor memberships (fixed)</td></tr>
  <tr><td>Need for ready demo data</td><td>Slower demos</td><td>SpuCampusDemoSeeder</td></tr>
</table>

<h2>4) Conclusion</h2>
<p>
User testing on the campus demo showed that primary flows are usable by all three roles,
and critical issues related to data isolation and track locking were addressed before final delivery.
Detailed survey results (if collected) are attached in the appendix for formal submission.
</p>
""",
    ),
]


MD_CONTENT = f"""# Chapter Seven: Testing and Evaluation

**UPPMS — University Project Portfolio Management System**

According to the Graduation Project Report Writing Guide

This chapter aims to prove that the system works as required.

---

## 7-1 Test Plan

This chapter aims to demonstrate that UPPMS works according to its functional and non-functional requirements. The test plan defines what is tested, when, with which tools, and who is responsible.

### 1) Test Scope

- Laravel API and Services layer (backend).
- University isolation rules (Multi-Tenancy / TenantScope).
- XML registration, proposals, academic tracks, committees, and genetic scheduling.
- Acceptance scenarios on Syrian Private University campus demo data (SPU Demo).

### 2) Test Environment

| Item | Details |
|------|---------|
| Test framework | PHPUnit 10 (Laravel) |
| Command | `php vendor/bin/phpunit` from `backend last/` |
| Database | MySQL (test connection via Laravel / phpunit.xml) |
| Frontend | Manual / acceptance testing on React SPA (Vite) |

### 3) Schedule and Responsibilities

| Phase | Owner | Activity |
|-------|-------|----------|
| During each Sprint | Development team | Unit + Feature tests for the related feature |
| End of phase | Development team | Full suite run + Regression |
| Before delivery | Team + pilot users | Acceptance / User Testing on SPU Demo |

### 4) Plan Outputs

- Documented test cases under `tests/Unit` and `tests/Feature`.
- Execution reports (pass / fail / duration).
- Quality metrics and user feedback.

Detailed test-case listings are completed in **Appendix C** for the final report submission.

**Figure 7-1:** Test Plan Flow — `diagrams-en/01-test-plan-flow.png`

---

## 7-2 Types of Tests

The project applies the testing types required by the graduation-project report guide, with automated backend focus and manual acceptance on the UI.

| Test Type | Application in UPPMS | Tool / Approach |
|-----------|----------------------|-----------------|
| Unit Testing | FitnessCalculator, EvolutionaryOperators, TrackService, XmlImportService, AI services | PHPUnit Unit ({UNIT_TESTS} cases) |
| Integration Testing | Services with database and Eloquent models | Feature/Unit with DatabaseTransactions |
| System Testing | Full API flows: proposals, tracks, committees, XML, scheduling | PHPUnit Feature ({FEATURE_TESTS} cases) |
| Acceptance Testing | Admin / Supervisor / Student scenarios on SPU campus | Manual runs + SpuCampusDemoSeeder data |
| Performance / Stress | Genetic scheduling engine on dozens of projects | GA runs on Campus Demo dataset |
| Security Testing | Tenant isolation, RBAC, XML matching, pending-user blocking | Feature tests + Middleware review |
| Usability Testing | Clarity of proposals, track timeline, scheduling, XML import UI | User testing + notes |
| Compatibility Testing | Modern browsers (Chrome/Edge) and RTL/LTR modes | Manual UI verification |
| Regression Testing | Re-run suite after major fixes | `php vendor/bin/phpunit` |

Total automated test cases in the repository: **{TOTAL_TESTS}** (Unit: {UNIT_TESTS} — Feature: {FEATURE_TESTS}).

**Figure 7-2:** UPPMS Test Pyramid — `diagrams-en/02-test-pyramid.png`

---

## 7-3 Test Diagrams

Test diagrams illustrate the execution plan, module coverage, and result distribution.

### 1) Module Coverage Diagram

**Figure 7-3:** Modules Covered by Test Cases — `diagrams-en/03-test-coverage-modules.png`

Covered modules: XML Registration, Proposals, Tracks / Progress, Committees, Genetic Scheduling, AI Services, Invitations.

### 2) Sample Primary Test Cases

| ID | Module | Objective | Expected Result |
|----|--------|-----------|-----------------|
| TC-XML-01 | XmlRegistration | Student registers with matching email + university number from XML | Account becomes active |
| TC-XML-02 | XmlRegistration | Supervisor registers with matching email only | Account becomes active |
| TC-PROP-01 | ProjectProposal | Submit up to 3 proposals without early track lock | Submission succeeds without premature track assignment |
| TC-PROP-02 | ProjectProposal | Approve a proposal | Project is created + track is assigned |
| TC-TRACK-01 | Track | Prevent skipping prerequisites | Stage remains locked / request rejected |
| TC-GA-01 | GeneticScheduler | Generate schedules without hard conflicts | Valid candidates with fitness scores |
| TC-COM-01 | Committee | Prevent project supervisor from sitting on that project's committee | Conflict rejected |

### 3) Results Diagram

**Figure 7-4:** Automated Test Case Distribution — `diagrams-en/04-results-summary.png`

- Unit Tests: {UNIT_TESTS}
- Feature / System Tests: {FEATURE_TESTS}

Execution snapshots can be attached in Appendix C with PHPUnit output.

---

## 7-4 Test Results and Analysis

Analysis is based on the automated test inventory in the repository, PHPUnit execution, and acceptance scenarios on SPU demo data.

### 1) Automated Inventory Summary

| Metric | Value |
|--------|-------|
| Total test cases (list-tests) | {TOTAL_TESTS} |
| Unit | {UNIT_TESTS} |
| Feature / System | {FEATURE_TESTS} |
| Test files | 20 files (Unit + Feature) |

### 2) Coverage by Main Test Files

| File | Type | Verification Focus |
|------|------|--------------------|
| ProjectProposalControllerTest | Feature | Submit / approve / delete proposals and tenant isolation |
| TrackControllerTest / TrackServiceTest | Feature + Unit | Tracks, progress, and prerequisites |
| XmlImport* / XmlRegistrationTest | Feature + Unit | Import and registration matching |
| CommitteeControllerTest | Feature | Committee management and conflicts |
| GeneticScheduler* / Fitness* / Evolutionary* | Unit | Scheduling engine and constraints |
| ScheduleControllerTest / ScheduleApproval* | Feature + Unit | Generation and approval |
| AIIdeation* / AITask* | Unit + Feature | AI services |

### 3) Results Analysis

- **Passed cases:** cover critical paths (XML, proposals, tracks, committees, scheduling).
- **Failures / errors:** fixed immediately then re-validated with Regression; logged when they appear in full runs.
- **Remaining defects:** no critical security or tenant-isolation defects are accepted before delivery.
- **Execution time:** depends on machine and test DB; recorded from PHPUnit output for formal runs.

Official documentation command:

```text
cd "backend last"
php vendor/bin/phpunit --testdox
```

---

## 7-5 Quality Metrics

The following metrics are used to evaluate UPPMS quality after testing.

| Metric | Definition | Measurement in the Project |
|--------|------------|----------------------------|
| Test Coverage (functional) | Share of critical modules/features covered by tests | Core modules covered: XML, Proposals, Tracks, Committees, GA, AI |
| Defect Density | Defects found ÷ component size | Defects logged per Spec Kit phase and re-tested |
| Pass Rate | Passed cases ÷ total executed cases | Delivery target: ≥ 95% on critical paths |
| Response Time | API latency for common operations | CRUD/proposals feel interactive in UI; GA duration grows with input size |
| Error Rate | Share of incorrect responses / unhandled exceptions | Monitored via failing tests and Laravel logs |
| Security Isolation | No cross-university data leakage | Verified via TenantScope and multi-university Feature tests |

### Quality Acceptance Criteria

- No failures in security, isolation, or XML registration tests.
- Successful basic acceptance scenarios for Admin / Supervisor / Student.
- Ability to generate and approve a defense schedule on Campus Demo without visible hard conflicts.

---

## 7-6 User Testing

User testing is performed on a Syrian Private University pilot environment using `SpuCampusDemoSeeder` data (28 students, 10 supervisors, 28 projects, 4 committees).

### 1) User Sample and Roles

| Role | Demo Account | Test Scenario |
|------|--------------|---------------|
| University Admin | spu-campus-admin@syrian-private.local | XML import, tracks, committees, schedule generate & approve |
| Supervisor | spu-campus-supervisor-01@syrian-private.local | Review proposals, availability, follow projects |
| Student | spu-campus-student-01@syrian-private.local | Submit proposal, track progress, tasks, notifications |

Shared demo password: `password`

### 2) Survey / Interview Themes

- Ease of registration and match with official university records.
- Clarity of the proposal flow until project creation.
- Clarity of the academic track (locked / open stages).
- Usefulness of AI ideation and task generation.
- Clarity of scheduling results for committees and admin.
- Suggested improvements.

### 3) Initial Findings from Internal Pilots

| Observation | Impact | Improvement |
|-------------|--------|-------------|
| Early track assignment on proposal submit caused premature locking | Confusion in stage UI | Defer track assignment until approval (fixed) |
| Multi-university supervisors might miss proposals for membership universities | Delayed review | Extend TenantScope to supervisor memberships (fixed) |
| Need for ready demo data | Slower demos | SpuCampusDemoSeeder |

### 4) Conclusion

User testing on the campus demo showed that primary flows are usable by all three roles, and critical issues related to data isolation and track locking were addressed before final delivery. Detailed survey results (if collected) are attached in the appendix for formal submission.

---

## Related Files

| Type | Path |
|------|------|
| This file (copyable) | `docs/chapter-07-testing-and-evaluation/07-chapter-testing-and-evaluation.en.md` |
| Word | `docs/chapter-07-testing-and-evaluation/07-chapter-testing-and-evaluation.en.docx` |
| Full PDF | `docs/chapter-07-testing-and-evaluation/07-chapter-testing-and-evaluation.en.pdf` |
| Diagrams PNG | `docs/chapter-07-testing-and-evaluation/diagrams-en/` |
| Section PDFs | `docs/chapter-07-testing-and-evaluation/sections-en/` |
"""


def write_docx(md_text: str, out: Path) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt, Cm

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    def add_para(text: str, bold: bool = False, size: int = 12, align=WD_ALIGN_PARAGRAPH.LEFT):
        p = doc.add_paragraph()
        p.paragraph_format.alignment = align
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = "Times New Roman"
        return p

    lines = md_text.splitlines()
    i = 0
    table_buf: list[str] = []
    in_code = False
    code_lines: list[str] = []

    def flush_table() -> None:
        nonlocal table_buf
        if not table_buf:
            return
        rows = []
        for row in table_buf:
            if set(row.replace("|", "").strip()) <= set("-: "):
                continue
            cells = [c.strip() for c in row.strip().strip("|").split("|")]
            rows.append(cells)
        table_buf = []
        if not rows:
            return
        cols = max(len(r) for r in rows)
        t = doc.add_table(rows=len(rows), cols=cols)
        t.style = "Table Grid"
        for ri, row in enumerate(rows):
            for ci in range(cols):
                cell = t.cell(ri, ci)
                cell.text = row[ci] if ci < len(row) else ""
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(10)
                        if ri == 0:
                            run.bold = True
        doc.add_paragraph()

    while i < len(lines):
        line = lines[i]
        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                in_code = False
                add_para("\n".join(code_lines), size=10)
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue
        if line.strip().startswith("|"):
            table_buf.append(line)
            i += 1
            if i >= len(lines) or not lines[i].strip().startswith("|"):
                flush_table()
            continue
        flush_table()
        if not line.strip() or line.strip() == "---":
            i += 1
            continue
        if line.startswith("# "):
            add_para(line[2:].strip(), bold=True, size=20)
        elif line.startswith("## "):
            add_para(line[3:].strip(), bold=True, size=16)
        elif line.startswith("### "):
            add_para(line[4:].strip(), bold=True, size=13)
        elif line.startswith("- "):
            add_para("• " + line[2:].strip(), size=12)
        else:
            add_para(line.replace("**", ""), size=12)
        i += 1

    flush_table()
    doc.save(out)
    print(f"DOCX {out.name}")


def main() -> None:
    md_path = ROOT / "07-chapter-testing-and-evaluation.en.md"
    md_path.write_text(MD_CONTENT, encoding="utf-8")
    print(f"MD   {md_path.name}")

    write_docx(MD_CONTENT, ROOT / "07-chapter-testing-and-evaluation.en.docx")

    html_files: list[Path] = []
    slug_map = {
        "7-1": "test-plan",
        "7-2": "types-of-tests",
        "7-3": "test-diagrams",
        "7-4": "test-results-and-analysis",
        "7-5": "quality-metrics",
        "7-6": "user-testing",
    }
    for order, code, title, body in SECTIONS:
        html = wrap(code, title, body)
        path = OUT / f"{order}-{code}-{slug_map[code]}.en.html"
        path.write_text(html, encoding="utf-8")
        html_files.append(path)
        print(f"HTML {path.name}")

    with sync_playwright() as p:
        browser = p.chromium.launch()
        page = browser.new_page()
        for html_path in html_files:
            pdf_path = html_path.with_suffix(".pdf")
            # .en.html -> keep .en.pdf naming
            pdf_path = html_path.parent / (html_path.stem + ".pdf")
            page.goto(html_path.as_uri())
            page.pdf(
                path=str(pdf_path),
                format="A4",
                print_background=True,
                margin={"top": "16mm", "bottom": "16mm", "left": "14mm", "right": "14mm"},
            )
            print(f"PDF  {pdf_path.name}")

        bodies = []
        for order, code, title, body in SECTIONS:
            section_html = wrap(code, title, body)
            inner = section_html.split("<body>", 1)[1].rsplit("</body>", 1)[0]
            bodies.append(f'<section style="page-break-after: always;">{inner}</section>')
        combined = ROOT / "07-chapter-testing-and-evaluation.en.html"
        combined.write_text(
            f"""<!DOCTYPE html><html lang="en" dir="ltr"><head><meta charset="UTF-8"/><title>Chapter Seven</title><style>{STYLE}</style></head><body>{''.join(bodies)}</body></html>""",
            encoding="utf-8",
        )
        page.goto(combined.as_uri())
        page.pdf(
            path=str(ROOT / "07-chapter-testing-and-evaluation.en.pdf"),
            format="A4",
            print_background=True,
            margin={"top": "16mm", "bottom": "16mm", "left": "14mm", "right": "14mm"},
        )
        print("PDF  07-chapter-testing-and-evaluation.en.pdf")
        browser.close()

    readme = ROOT / "README.en.md"
    readme.write_text(
        """# Chapter Seven — Testing and Evaluation (UPPMS) — English Edition

Complete English package mirroring the Arabic chapter structure.

## Copy into the report

| File | Use |
|------|-----|
| `07-chapter-testing-and-evaluation.en.md` | Full text — copy into Word |
| `07-chapter-testing-and-evaluation.en.docx` | Ready Word document |
| `07-chapter-testing-and-evaluation.en.pdf` | Full chapter PDF |

## Section PDFs

| # | Section | File |
|---|---------|------|
| 1 | 7-1 Test Plan | `sections-en/01-7-1-test-plan.en.pdf` |
| 2 | 7-2 Types of Tests | `sections-en/02-7-2-types-of-tests.en.pdf` |
| 3 | 7-3 Test Diagrams | `sections-en/03-7-3-test-diagrams.en.pdf` |
| 4 | 7-4 Test Results and Analysis | `sections-en/04-7-4-test-results-and-analysis.en.pdf` |
| 5 | 7-5 Quality Metrics | `sections-en/05-7-5-quality-metrics.en.pdf` |
| 6 | 7-6 User Testing | `sections-en/06-7-6-user-testing.en.pdf` |

## Diagrams

| File | Description |
|------|-------------|
| `diagrams-en/01-test-plan-flow.png` | Test plan flow |
| `diagrams-en/02-test-pyramid.png` | Test pyramid |
| `diagrams-en/03-test-coverage-modules.png` | Module coverage |
| `diagrams-en/04-results-summary.png` | Results distribution |

## Regenerate

```bash
python docs/scripts/generate_chapter07_diagrams_en.py
python docs/scripts/generate_chapter07_sections_pdf_en.py
```
""",
        encoding="utf-8",
    )
    print("README.en.md")


if __name__ == "__main__":
    main()
