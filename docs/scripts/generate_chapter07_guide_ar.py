#!/usr/bin/env python3
"""
Chapter 7 — Testing & Evaluation
Formatted exactly per SPU guide (د. كادان الجمعة):
- Arabic: Simplified Arabic (18/16/14/14), line-spacing 1.5, margins 2.54cm
- Tables: caption ABOVE — «جدول س. العنوان»
- Figures: caption BELOW — «شكل س. العنوان»
- Sections: 7-1 … 7-6 + Appendix C (ملحق ج)
"""
from __future__ import annotations

from pathlib import Path

from playwright.sync_api import sync_playwright

ROOT = Path(__file__).resolve().parents[1] / "chapter-07-testing-and-evaluation"
OUT = ROOT / "sections"
DIAG = ROOT / "diagrams"
OUT.mkdir(parents=True, exist_ok=True)

TOTAL, UNIT, FEATURE = 141, 50, 91

# Guide: margins 2.54cm, line-spacing 1.5
STYLE_AR = """
  @page { size: A4; margin: 2.54cm; }
  body {
    font-family: "Simplified Arabic", "Traditional Arabic", "Arial", sans-serif;
    font-size: 14pt; line-height: 1.5; color: #000; margin: 0; background: #fff;
    text-align: justify; direction: rtl;
  }
  h1.chapter {
    font-size: 18pt; font-weight: bold; text-align: center;
    margin: 0 0 22px; page-break-after: avoid;
  }
  h2 {
    font-size: 16pt; font-weight: bold; margin: 20px 0 10px;
    page-break-after: avoid; text-align: right;
  }
  h3 {
    font-size: 14pt; font-weight: bold; margin: 14px 0 8px;
    page-break-after: avoid; text-align: right;
  }
  p { margin: 0 0 10px; text-indent: 1.2em; }
  p.ni { text-indent: 0; }
  ul, ol { margin: 6px 0 12px; padding-right: 1.5em; }
  li { margin: 4px 0; }
  table {
    width: 100%; border-collapse: collapse; margin: 6px 0 14px;
    font-size: 12pt; direction: rtl; page-break-inside: avoid;
  }
  th, td {
    border: 1px solid #000; padding: 6px 8px;
    text-align: right; vertical-align: top;
  }
  th { background: #f0f0f0; font-weight: bold; }
  .tbl-cap {
    font-size: 14pt; font-weight: bold; text-align: center;
    margin: 14px 0 6px; text-indent: 0;
  }
  .fig-cap {
    font-size: 14pt; font-weight: bold; text-align: center;
    margin: 6px 0 14px; text-indent: 0;
  }
  .imgbox { margin: 10px 0 4px; text-align: center; page-break-inside: avoid; }
  .imgbox img { max-width: 95%; height: auto; border: 1px solid #999; }
  .note {
    margin: 12px 0; padding: 8px 12px; border: 1px solid #555;
    text-indent: 0; font-size: 13pt;
  }
  .page-break { page-break-before: always; }
"""


def table(caption: str, headers: list[str], rows: list[list[str]]) -> str:
    head = "".join(f"<th>{h}</th>" for h in headers)
    body = "".join(
        "<tr>" + "".join(f"<td>{c}</td>" for c in row) + "</tr>" for row in rows
    )
    return f"""
<p class="tbl-cap">{caption}</p>
<table>
  <tr>{head}</tr>
  {body}
</table>
"""


def figure(path: Path, caption: str) -> str:
    if path.exists():
        return f"""
<div class="imgbox"><img src="{path.as_uri()}" alt="{caption}" /></div>
<p class="fig-cap">{caption}</p>
"""
    return f'<p class="note">[يُدرج الشكل: {caption}]</p>'


def wrap(title: str, body: str) -> str:
    return f"""<!DOCTYPE html>
<html lang="ar" dir="rtl">
<head><meta charset="UTF-8"/><title>{title}</title>
<style>{STYLE_AR}</style></head>
<body>
{body}
</body></html>
"""


CHAPTER_BODY = f"""
<h1 class="chapter">الفصل السابع: الاختبار والتقييم</h1>

<p>
يهدف هذا الفصل إلى إثبات أن نظام إدارة محافظ مشاريع الجامعات
(UPPMS — University Project Portfolio Management System)
يعمل كما هو مطلوب وفق المتطلبات الوظيفية وغير الوظيفية.
ولذلك يُعرض فيما يأتي خطة الاختبار، وأنواع الاختبارات المنفَّذة، ومخططاتها،
ثم نتائجها وتحليلها، ومقاييس الجودة، وأخيراً نتائج اختبار المستخدمين.
ويُكمَّل التفصيل التشغيلي لحالات الاختبار في الملحق ج.
</p>

<h2>7-1 خطة الاختبار</h2>

<p>
تُحدَّد خطة الاختبار بالإجابة عن الأسئلة الآتية: ماذا سنختبر؟ ومتى؟ وبأي أدوات؟
ومن المسؤول عن الاختبار؟ ويوضح الشكل الآتي تدفق تنفيذ الخطة من تحديد النطاق
حتى تحليل التغذية الراجعة.
</p>

{figure(DIAG / "01-test-plan-flow.png", "شكل 7-1. مخطط تدفق خطة الاختبار")}

<p>
يبين الجدول الآتي نطاق الاختبار في نظام UPPMS، وما يشمله وما يُستثنى منه في هذه المرحلة.
</p>

{table(
    "جدول 7-1. نطاق الاختبار",
    ["البند", "التفاصيل"],
    [
        ["ما سيتم اختباره", "وحدات الباكند (Services)، تكاملها مع قاعدة البيانات، مسارات واجهة برمجة التطبيقات (API)، عزل الجامعات، التسجيل عبر XML، المقترحات، المسارات الأكاديمية، اللجان، الجدولة الجينية، وسيناريوهات القبول على بيانات الحرم التجريبي"],
        ["ما لن يُختبر بالكامل هنا", "اختبار توافق شامل لكل المتصفحات والأجهزة، واختبارات تحمل إنتاجية واسعة النطاق خارج بيئة التطوير"],
        ["متى", "أثناء كل Sprint للميزات الجديدة، وفي نهاية كل مرحلة للمجموعة الكاملة، وقبل التسليم لاختبارات القبول"],
        ["الأدوات", "PHPUnit 10 ضمن Laravel، وMySQL لبيئة الاختبار، مع تحقق يدوي على واجهة React"],
        ["المسؤول", "فريق التطوير (تنفيذ آلي وتحليل)، والمشرف الأكاديمي (مراجعة)، ومستخدمو التجربة (اختبار قبول)"],
    ],
)}

<p>
أما الجدول الآتي فيوضح توزيع المسؤوليات خلال دورة حياة الاختبار.
</p>

{table(
    "جدول 7-2. مسؤوليات الاختبار والجدول الزمني",
    ["المرحلة", "المسؤول", "النشاط"],
    [
        ["أثناء كل Sprint", "فريق التطوير", "كتابة وتشغيل اختبارات الوحدة والنظام المرتبطة بالميزة"],
        ["نهاية المرحلة", "فريق التطوير", "تشغيل المجموعة الكاملة وإجراء اختبارات الانحدار"],
        ["قبل التسليم", "الفريق + مستخدمون تجريبيون", "اختبارات القبول واختبار المستخدمين على بيانات SPU"],
    ],
)}

<p class="ni note">
يُراجع الملحق ج لنموذج خطة الاختبار التفصيلي وحالات الاختبار الموسَّعة.
</p>

<h2>7-2 أنواع الاختبارات</h2>

<p>
اعتمد المشروع أنواعاً متعددة من الاختبارات وفق ما نص عليه الدليل، وهي:
اختبارات الوحدة، والتكامل، والنظام، والقبول، والأداء والتحمل، والأمان،
وقابلية الاستخدام، والتوافق، والانحدار. ويبين الجدول الآتي كيفية تطبيق كل نوع
على نظام UPPMS.
</p>

{table(
    "جدول 7-3. أنواع الاختبارات وتطبيقها في UPPMS",
    ["نوع الاختبار", "الهدف في المشروع", "الأداة / الأسلوب"],
    [
        ["اختبارات الوحدة (Unit)", "التحقق من دوال/وحدات منفصلة مثل حساب اللياقة ومشغّلات الخوارزمية الجينية وخدمات المسارات وXML", f"PHPUnit Unit ({UNIT} حالة)"],
        ["اختبارات التكامل (Integration)", "التحقق من تفاعل الخدمات مع قاعدة البيانات ونماذج Eloquent", "Feature/Unit مع DatabaseTransactions"],
        ["اختبارات النظام (System)", "التحقق من مسارات API كاملة: المقترحات، المسارات، اللجان، XML، الجدولة", f"PHPUnit Feature ({FEATURE} حالة)"],
        ["اختبارات القبول (Acceptance)", "التحقق من سيناريوهات أدمن الجامعة والمشرف والطالب", "تجربة يدوية + SpuCampusDemoSeeder"],
        ["الأداء والتحمل", "قياس سلوك محرك الجدولة الجينية مع عشرات المشاريع", "تشغيل الخوارزمية على بيانات الحرم التجريبي"],
        ["الأمان", "عزل الجامعات، الصلاحيات، مطابقة XML، منع وصول الحسابات المعلّقة", "Feature tests + مراجعة Middleware"],
        ["قابلية الاستخدام", "وضوح واجهات المقترحات والمسار والجدولة والاستيراد", "اختبار مستخدمين + ملاحظات"],
        ["التوافق", "المتصفحات الحديثة ووضعا RTL/LTR", "تحقق يدوي على الواجهة"],
        ["الانحدار (Regression)", "التأكد من عدم كسر الميزات السابقة بعد الإصلاحات", "php vendor/bin/phpunit"],
    ],
)}

<p>
ويوضح الشكل الآتي هرم الاختبارات المعتمد في المشروع، حيث تشكل اختبارات الوحدة القاعدة الأوسع،
تليها اختبارات التكامل والنظام، ثم اختبارات القبول في القمة.
</p>

{figure(DIAG / "02-test-pyramid.png", "شكل 7-2. هرم أنواع الاختبارات في UPPMS")}

<p>
بلغ إجمالي حالات الاختبار الآلية المُدرَجة في المستودع
<strong>{TOTAL}</strong> حالة، منها {UNIT} لاختبارات الوحدة و{FEATURE} لاختبارات النظام/الميزات.
</p>

<h2>7-3 مخططات الاختبار</h2>

<p>
تشمل مخططات الاختبار مخططات حالات الاختبار
(Test Case Diagrams)
ومخططات نتائج الاختبار. ويبين الشكل الآتي الوحدات البرمجية المغطاة بحالات الاختبار الآلية.
</p>

{figure(DIAG / "03-test-coverage-modules.png", "شكل 7-3. مخطط تغطية وحدات النظام بالاختبارات")}

<p>
يعرض الجدول الآتي نماذج من حالات الاختبار الرئيسية (ويُستكمل التفصيل في الملحق ج).
</p>

{table(
    "جدول 7-4. نماذج حالات الاختبار (Test Cases)",
    ["رقم الحالة", "الوصف", "المدخلات", "الخرج المتوقع"],
    [
        ["TC-XML-01", "تسجيل طالب مطابق لسجل XML", "إيميل + رقم جامعي مطابقان", "حساب بحالة active"],
        ["TC-XML-02", "تسجيل مشرف مطابق لسجل XML", "إيميل مطابق فقط", "حساب بحالة active"],
        ["TC-PROP-01", "تقديم مقترحات دون قفل مبكر للمسار", "حتى 3 مقترحات معلقة", "نجاح التقديم دون تعيين مسار مبكر"],
        ["TC-PROP-02", "اعتماد مقترح مشروع", "موافقة المشرف", "إنشاء مشروع وتعيين المسار"],
        ["TC-TRACK-01", "منع تخطي المتطلبات السابقة", "محاولة فتح مرحلة لاحقة دون نجاح السابقة", "رفض / إبقاء المرحلة مقفلة"],
        ["TC-GA-01", "توليد جداول مناقشة", "مشاريع + توافر + قاعات", "مرشحون صالحون مع قيمة fitness"],
        ["TC-COM-01", "منع تعارض مصلحة في اللجنة", "مشرف المشروع عضواً في لجنة مناقشته", "رفض التعارض"],
    ],
)}

<p>
أما توزيع حالات الاختبار الآلية بين اختبارات الوحدة واختبارات النظام فيوضحه الشكل الآتي.
</p>

{figure(DIAG / "04-results-summary.png", "شكل 7-4. مخطط نتائج / توزيع حالات الاختبار الآلية")}

<h2>7-4 نتائج الاختبار وتحليلها</h2>

<p>
يعتمد تحليل النتائج على مخزون حالات الاختبار الآلية وتشغيل PHPUnit،
إضافة إلى سيناريوهات القبول على بيانات الجامعة الخاصة السورية.
ويبين الجدول الآتي ملخص المخزون الآلي.
</p>

{table(
    "جدول 7-5. ملخص مخزون الاختبارات الآلية",
    ["المؤشر", "القيمة"],
    [
        ["إجمالي حالات الاختبار", str(TOTAL)],
        ["اختبارات الوحدة (Unit)", str(UNIT)],
        ["اختبارات النظام/الميزات (Feature)", str(FEATURE)],
        ["عدد ملفات الاختبار", "20 ملفاً"],
    ],
)}

<p>
ويوضح الجدول الآتي توزيع التغطية على الملفات الرئيسية في المشروع.
</p>

{table(
    "جدول 7-6. توزيع التغطية حسب ملفات الاختبار الرئيسية",
    ["الملف", "النوع", "مجال التحقق"],
    [
        ["ProjectProposalControllerTest", "Feature", "تقديم/اعتماد/حذف المقترح وعزل الجامعات"],
        ["TrackControllerTest / TrackServiceTest", "Feature + Unit", "المسارات والتقدم والمتطلبات السابقة"],
        ["XmlImport* / XmlRegistrationTest", "Feature + Unit", "الاستيراد والمطابقة عند التسجيل"],
        ["CommitteeControllerTest", "Feature", "إدارة اللجان والتعارضات"],
        ["GeneticScheduler* / Fitness* / Evolutionary*", "Unit", "محرك الجدولة والقيود"],
        ["ScheduleControllerTest / ScheduleApproval*", "Feature + Unit", "التوليد والاعتماد"],
        ["AIIdeation* / AITask*", "Unit + Feature", "خدمات الذكاء الاصطناعي"],
    ],
)}

<p>
ومن تحليل النتائج يمكن استخلاص ما يأتي:
</p>
<ul>
  <li>تغطي حالات النجاح المسارات الحرجة: XML، والمقترحات، والمسارات، واللجان، والجدولة.</li>
  <li>تُعالَج حالات الفشل فور اكتشافها، ثم يُعاد تشغيل اختبارات الانحدار للتحقق من الإصلاح.</li>
  <li>لا تُقبل أخطاء متبقية حرجة على مسارات الأمن أو عزل الجامعات قبل التسليم.</li>
  <li>يُسجَّل زمن تنفيذ الاختبارات من مخرجات PHPUnit عند كل تشغيل رسمي للتوثيق.</li>
</ul>

<p class="ni note">
أمر التشغيل الرسمي: من مجلد <code>backend last</code> يُنفَّذ الأمر
<code>php vendor/bin/phpunit --testdox</code>.
</p>

<h2>7-5 مقاييس الجودة</h2>

<p>
تُقاس جودة النظام بعد الاختبار باستخدام مجموعة من المقاييس التي نص عليها الدليل،
وهي: تغطية الاختبارات (Test Coverage)، وكثافة العيوب (Defect Density)،
وزمن الاستجابة، ومعدل الأخطاء. ويبين الجدول الآتي تعريف كل مقياس وكيفية تطبيقه في المشروع.
</p>

{table(
    "جدول 7-7. مقاييس الجودة المستخدمة في UPPMS",
    ["المقياس", "التعريف", "القياس في المشروع"],
    [
        ["تغطية الاختبارات (Test Coverage)", "نسبة الوحدات/الميزات الحرجة المغطاة باختبارات", "تغطية الوحدات الأساسية: XML، المقترحات، المسارات، اللجان، الجدولة الجينية، والذكاء الاصطناعي"],
        ["كثافة العيوب (Defect Density)", "عدد العيوب المكتشفة ÷ حجم المكوّن", "تُسجَّل العيوب لكل مرحلة تطويرية (Spec Kit) ويُعاد اختبارها بعد الإصلاح"],
        ["معدل النجاح (Pass Rate)", "عدد الحالات الناجحة ÷ إجمالي الحالات المنفَّذة", "مستهدف التسليم: لا يقل عن 95٪ للمسارات الحرجة"],
        ["زمن الاستجابة", "الزمن اللازم لاستجابة واجهة برمجة التطبيقات للعمليات الاعتيادية", "عمليات الإدارة والمقترحات ضمن تفاعل فوري؛ الجدولة الجينية أطول بحسب حجم المدخلات"],
        ["معدل الأخطاء", "نسبة الاستجابات الخاطئة أو الاستثناءات غير المعالجة", "تُراقب عبر فشل الاختبارات وسجلات Laravel"],
        ["عزل الجامعات", "عدم تسرّب البيانات بين الجامعات", "مُتحقَّق عبر TenantScope واختبارات متعددة الجامعات"],
    ],
)}

<p>
وبناءً على هذه المقاييس، اعتمد المشروع معايير قبول للجودة تشمل:
عدم وجود فشل في اختبارات الأمن والعزل والتسجيل عبر XML،
ونجاح سيناريوهات القبول للأدوار الثلاثة (مدير الجامعة، المشرف، الطالب)،
وإمكانية توليد واعتماد جدول مناقشة على بيانات الحرم التجريبي دون تعارضات صارمة ظاهرة.
</p>

<h2>7-6 اختبار المستخدمين</h2>

<p>
أُجري اختبار المستخدمين على بيئة تجريبية لجامعة
<strong>syrian private uni</strong>
باستخدام بيانات الحرم التجريبي
(<code>SpuCampusDemoSeeder</code>)
التي تضم 28 طالباً و10 مشرفين و28 مشروعاً و4 لجان.
ويبين الجدول الآتي عينة المستخدمين وسيناريوهات الاختبار.
</p>

{table(
    "جدول 7-8. عينة اختبار المستخدمين وسيناريوهاتهم",
    ["الدور", "الحساب التجريبي", "سيناريو الاختبار"],
    [
        ["مدير الجامعة", "spu-campus-admin@syrian-private.local", "استيراد XML، إدارة المسارات واللجان، توليد الجدول واعتماده"],
        ["مشرف", "spu-campus-supervisor-01@syrian-private.local", "مراجعة المقترحات، تسجيل أوقات الفراغ، متابعة المشاريع"],
        ["طالب", "spu-campus-student-01@syrian-private.local", "تقديم مقترح، متابعة التقدم الأكاديمي، المهام، والإشعارات"],
    ],
)}

<p class="ni">كلمة المرور التجريبية الموحدة لجميع الحسابات: <code>password</code>.</p>

<p>
وركّزت محاور الاستبيان/المقابلات على: سهولة التسجيل وتوافق البيانات مع سجلات الجامعة،
ووضوح مسار المقترح حتى إنشاء المشروع، ووضوح المسار الأكاديمي،
وفائدة اقتراحات الذكاء الاصطناعي، ووضوح نتائج الجدولة، والتحسينات المقترحة.
</p>

<p>
ويعرض الجدول الآتي أبرز الملاحظات الناتجة عن التجريب الداخلي والتحسينات المنفَّذة.
</p>

{table(
    "جدول 7-9. ملاحظات المستخدمين والتحسينات المقترحة",
    ["ملاحظة المستخدم", "الأثر", "التحسين المنفَّذ / المقترح"],
    [
        ["الربط المبكر بالمسار عند تقديم المقترح", "التباس في حالة المراحل الأكاديمية", "تأجيل تعيين المسار حتى اعتماد المقترح (تم تنفيذه)"],
        ["المشرف المرتبط بأكثر من جامعة قد لا يرى مقترحات جامعة العضوية", "تأخر مراجعة المقترحات", "توسيع نطاق TenantScope لعضويات المشرف (تم تنفيذه)"],
        ["الحاجة إلى بيانات جاهزة للعرض والتجريب", "بطء تجهيز الديمو", "اعتماد SpuCampusDemoSeeder لبيانات حرم SPU"],
    ],
)}

<p>
وخلاصة القول إن اختبار المستخدمين أظهر قابلية استخدام المسارات الأساسية من قبل الأدوار الثلاثة،
وأن الملاحظات الحرجة المتعلقة بعزل البيانات وقفل المسار قد عُولجت قبل مرحلة التسليم النهائية.
وتُرفَق نتائج الاستبيان التفصيلية — إن وُجدت — ضمن ملاحق التقرير عند التسليم الرسمي.
</p>
"""

APPENDIX_C = f"""
<h1 class="chapter">ملحق ج — نموذج خطة الاختبار (UPPMS)</h1>

<p class="ni">
يتبع هذا الملحق نموذج الملحق ج الوارد في دليل كتابة تقارير مشاريع التخرج،
مكيَّفاً على نظام UPPMS.
</p>

<h2>ج-1 نطاق الاختبار</h2>
<ul>
  <li><strong>ما سيتم اختباره:</strong> الوحدات البرمجية للباكند، والتكامل بينها، وواجهة برمجة التطبيقات، وعزل الجامعات، والأداء الخاص بالجدولة الجينية، والأمان.</li>
  <li><strong>ما لن يتم اختباره بالكامل:</strong> الاختبارات اليدوية الشاملة لكل الشاشات، واختبارات التوافق مع جميع المتصفحات والأجهزة.</li>
</ul>

<h2>ج-2 استراتيجية الاختبار</h2>
{table(
    "جدول ج-1. استراتيجية الاختبار",
    ["نوع الاختبار", "الهدف", "الأداة"],
    [
        ["اختبارات الوحدة", "اختبار كل دالة/وحدة بشكل منفصل", "PHPUnit (Unit)"],
        ["اختبارات التكامل", "اختبار تفاعل الوحدات مع قاعدة البيانات", "PHPUnit + DatabaseTransactions"],
        ["اختبارات النظام", "اختبار مسارات API كاملة", "PHPUnit (Feature)"],
        ["اختبارات الأداء", "قياس سلوك الجدولة الجينية مع بيانات كثيفة", "تشغيل GeneticScheduler على Campus Demo"],
        ["اختبارات الأمان", "كشف مشاكل العزل والصلاحيات ومطابقة XML", "Feature tests + مراجعة Middleware"],
    ],
)}

<h2>ج-3 بيئة الاختبار</h2>
{table(
    "جدول ج-2. بيئة الاختبار",
    ["العنصر", "المواصفات"],
    [
        ["نظام التشغيل", "Windows 10/11 (بيئة التطوير المحلية)"],
        ["الخادم", "خادم تطوير محلي (localhost) — Laravel"],
        ["قاعدة البيانات", "MySQL (بيئة اختبار منفصلة عبر إعدادات phpunit.xml)"],
        ["الواجهة", "React SPA (Vite) على المتصفح"],
        ["المتصفحات", "Chrome / Edge"],
    ],
)}

<h2>ج-4 حالات الاختبار (نموذج)</h2>
{table(
    "جدول ج-3. نموذج حالات الاختبار",
    ["رقم الحالة", "الوصف", "المدخلات", "الخرج المتوقع", "النتيجة"],
    [
        ["TC-01", "تسجيل طالب مطابق لـ XML", "إيميل ورقم جامعي صحيحان من XML", "حساب active", "تم"],
        ["TC-02", "تسجيل طالب غير مطابق", "إيميل أو رقم غير موجود في XML", "رفض مع رسالة عدم التطابق", "تم"],
        ["TC-03", "تسجيل مشرف مطابق", "إيميل مطابق فقط", "حساب active", "تم"],
        ["TC-04", "تقديم مقترح مشروع", "عنوان ووصف ومشرف", "مقترح بحالة pending", "تم"],
        ["TC-05", "اعتماد المقترح", "موافقة المشرف", "إنشاء مشروع", "تم"],
        ["TC-06", "منع تخطي مرحلة أكاديمية", "محاولة مرحلة لاحقة دون نجاح السابقة", "رفض/قفل", "تم"],
        ["TC-07", "توليد جدول مناقشة", "مشاريع وتوافر وقاعات", "مرشحون للجداول", "تم"],
        ["TC-08", "منع تعارض اللجنة", "مشرف المشروع ضمن اللجنة", "رفض التعارض", "تم"],
    ],
)}
"""


def md_from_guide() -> str:
    return f"""# الفصل السابع: الاختبار والتقييم

يهدف هذا الفصل إلى إثبات أن نظام إدارة محافظ مشاريع الجامعات (UPPMS) يعمل كما هو مطلوب وفق المتطلبات الوظيفية وغير الوظيفية. ولذلك يُعرض فيما يأتي خطة الاختبار، وأنواع الاختبارات المنفَّذة، ومخططاتها، ثم نتائجها وتحليلها، ومقاييس الجودة، وأخيراً نتائج اختبار المستخدمين. ويُكمَّل التفصيل التشغيلي لحالات الاختبار في الملحق ج.

## 7-1 خطة الاختبار

تُحدَّد خطة الاختبار بالإجابة عن الأسئلة الآتية: ماذا سنختبر؟ ومتى؟ وبأي أدوات؟ ومن المسؤول عن الاختبار؟ (انظر الملحق ج). ويوضح شكل 7-1 تدفق تنفيذ الخطة.

**شكل 7-1.** مخطط تدفق خطة الاختبار  
`diagrams/01-test-plan-flow.png`

يبين جدول 7-1 نطاق الاختبار في نظام UPPMS.

**جدول 7-1.** نطاق الاختبار

| البند | التفاصيل |
|-------|----------|
| ما سيتم اختباره | وحدات الباكند (Services)، تكاملها مع قاعدة البيانات، مسارات API، عزل الجامعات، التسجيل عبر XML، المقترحات، المسارات الأكاديمية، اللجان، الجدولة الجينية، وسيناريوهات القبول على بيانات الحرم التجريبي |
| ما لن يُختبر بالكامل هنا | اختبار توافق شامل لكل المتصفحات، واختبارات تحمل إنتاجية واسعة خارج بيئة التطوير |
| متى | أثناء كل Sprint، ونهاية كل مرحلة، وقبل التسليم |
| الأدوات | PHPUnit 10 ضمن Laravel، وMySQL، مع تحقق يدوي على React |
| المسؤول | فريق التطوير، والمشرف الأكاديمي، ومستخدمو التجربة |

**جدول 7-2.** مسؤوليات الاختبار والجدول الزمني

| المرحلة | المسؤول | النشاط |
|---------|---------|--------|
| أثناء كل Sprint | فريق التطوير | اختبارات الوحدة والنظام المرتبطة بالميزة |
| نهاية المرحلة | فريق التطوير | المجموعة الكاملة + الانحدار |
| قبل التسليم | الفريق + مستخدمون تجريبيون | القبول واختبار المستخدمين على SPU |

## 7-2 أنواع الاختبارات

اعتمد المشروع: اختبارات الوحدة، والتكامل، والنظام، والقبول، والأداء والتحمل، والأمان، وقابلية الاستخدام، والتوافق، والانحدار.

**جدول 7-3.** أنواع الاختبارات وتطبيقها في UPPMS

| نوع الاختبار | الهدف في المشروع | الأداة / الأسلوب |
|--------------|------------------|------------------|
| الوحدة | دوال اللياقة ومشغّلات GA وخدمات المسارات وXML | PHPUnit Unit ({UNIT} حالة) |
| التكامل | تفاعل الخدمات مع قاعدة البيانات | Feature/Unit + DatabaseTransactions |
| النظام | مسارات API كاملة | PHPUnit Feature ({FEATURE} حالة) |
| القبول | سيناريوهات أدمن/مشرف/طالب | يدوي + SpuCampusDemoSeeder |
| الأداء والتحمل | سلوك الجدولة الجينية | تشغيل GA على Campus Demo |
| الأمان | العزل والصلاحيات وXML | Feature + Middleware |
| قابلية الاستخدام | وضوح الواجهات | اختبار مستخدمين |
| التوافق | Chrome/Edge وRTL/LTR | تحقق يدوي |
| الانحدار | عدم كسر الميزات السابقة | phpunit |

**شكل 7-2.** هرم أنواع الاختبارات في UPPMS  
`diagrams/02-test-pyramid.png`

إجمالي الحالات الآلية: **{TOTAL}** (Unit: {UNIT} — Feature: {FEATURE}).

## 7-3 مخططات الاختبار

تشمل مخططات الاختبار: Test Case Diagrams ومخططات نتائج الاختبار.

**شكل 7-3.** مخطط تغطية وحدات النظام بالاختبارات  
`diagrams/03-test-coverage-modules.png`

**جدول 7-4.** نماذج حالات الاختبار (Test Cases)

| رقم الحالة | الوصف | المدخلات | الخرج المتوقع |
|------------|-------|----------|---------------|
| TC-XML-01 | تسجيل طالب مطابق لـ XML | إيميل + رقم | حساب active |
| TC-XML-02 | تسجيل مشرف مطابق | إيميل فقط | حساب active |
| TC-PROP-01 | تقديم مقترحات دون قفل مبكر | حتى 3 مقترحات | نجاح دون تعيين مسار مبكر |
| TC-PROP-02 | اعتماد مقترح | موافقة المشرف | إنشاء مشروع + مسار |
| TC-TRACK-01 | منع تخطي المتطلبات | مرحلة لاحقة دون نجاح السابقة | رفض/قفل |
| TC-GA-01 | توليد جداول مناقشة | مشاريع+توافر+قاعات | مرشحون + fitness |
| TC-COM-01 | منع تعارض اللجنة | مشرف المشروع في اللجنة | رفض |

**شكل 7-4.** مخطط نتائج / توزيع حالات الاختبار الآلية  
`diagrams/04-results-summary.png`

## 7-4 نتائج الاختبار وتحليلها

**جدول 7-5.** ملخص مخزون الاختبارات الآلية

| المؤشر | القيمة |
|--------|--------|
| إجمالي حالات الاختبار | {TOTAL} |
| Unit | {UNIT} |
| Feature | {FEATURE} |
| ملفات الاختبار | 20 |

**جدول 7-6.** توزيع التغطية حسب ملفات الاختبار الرئيسية

| الملف | النوع | مجال التحقق |
|-------|-------|-------------|
| ProjectProposalControllerTest | Feature | مقترحات + عزل |
| TrackControllerTest / TrackServiceTest | Feature + Unit | مسارات وتقدم |
| XmlImport* / XmlRegistrationTest | Feature + Unit | XML |
| CommitteeControllerTest | Feature | لجان |
| GeneticScheduler* / Fitness* / Evolutionary* | Unit | جدولة |
| ScheduleControllerTest / ScheduleApproval* | Feature + Unit | توليد واعتماد |
| AIIdeation* / AITask* | Unit + Feature | ذكاء اصطناعي |

تحليل النتائج:
- تغطي حالات النجاح المسارات الحرجة.
- تُعالَج حالات الفشل ثم يُعاد الانحدار.
- لا تُقبل أخطاء حرجة على الأمن أو العزل قبل التسليم.
- يُسجَّل زمن التنفيذ من مخرجات PHPUnit.

## 7-5 مقاييس الجودة

**جدول 7-7.** مقاييس الجودة المستخدمة في UPPMS

| المقياس | التعريف | القياس في المشروع |
|---------|---------|-------------------|
| Test Coverage | نسبة الميزات الحرجة المغطاة | XML، مقترحات، مسارات، لجان، GA، AI |
| Defect Density | عيوب ÷ حجم المكوّن | تسجيل لكل مرحلة Spec Kit |
| Pass Rate | نجاح ÷ إجمالي | مستهدف ≥ 95٪ للمسارات الحرجة |
| زمن الاستجابة | زمن API | CRUD فوري؛ GA أطول حسب الحجم |
| معدل الأخطاء | استجابات خاطئة / استثناءات | عبر الاختبارات والسجلات |
| عزل الجامعات | عدم تسرّب البيانات | TenantScope + اختبارات متعددة الجامعات |

معايير القبول: لا فشل أمني/عزل/XML؛ نجاح قبول الأدوار الثلاثة؛ توليد واعتماد جدول مناقشة دون تعارضات صارمة ظاهرة.

## 7-6 اختبار المستخدمين

أُجري على بيانات `SpuCampusDemoSeeder` (28 طالباً، 10 مشرفين، 28 مشروعاً، 4 لجان).

**جدول 7-8.** عينة اختبار المستخدمين وسيناريوهاتهم

| الدور | الحساب التجريبي | سيناريو الاختبار |
|-------|-----------------|------------------|
| مدير الجامعة | spu-campus-admin@syrian-private.local | XML، مسارات، لجان، جدولة |
| مشرف | spu-campus-supervisor-01@syrian-private.local | مقترحات، توافر، مشاريع |
| طالب | spu-campus-student-01@syrian-private.local | مقترح، تقدم، مهام، إشعارات |

كلمة المرور: `password`

**جدول 7-9.** ملاحظات المستخدمين والتحسينات المقترحة

| ملاحظة المستخدم | الأثر | التحسين |
|-----------------|-------|---------|
| ربط مبكر بالمسار عند التقديم | التباس المراحل | تأجيل التعيين حتى الاعتماد (تم) |
| مشرف متعدد الجامعات لا يرى مقترحات العضوية | تأخر المراجعة | توسيع TenantScope (تم) |
| الحاجة لبيانات جاهزة | بطء الديمو | SpuCampusDemoSeeder |

الخلاصة: المسارات الأساسية قابلة للاستخدام من الأدوار الثلاثة، وعولجت الملاحظات الحرجة قبل التسليم.

---

# ملحق ج — نموذج خطة الاختبار (UPPMS)

## ج-1 نطاق الاختبار
- ما سيتم اختباره: الوحدات، التكامل، API، العزل، أداء الجدولة، الأمان.
- ما لن يُختبر بالكامل: توافق كل المتصفحات، وتحمل إنتاجي واسع.

## ج-2 استراتيجية الاختبار

**جدول ج-1.** استراتيجية الاختبار

| نوع الاختبار | الهدف | الأداة |
|--------------|-------|--------|
| الوحدة | كل دالة/وحدة منفصلة | PHPUnit Unit |
| التكامل | تفاعل الوحدات مع DB | PHPUnit + Transactions |
| النظام | مسارات API كاملة | PHPUnit Feature |
| الأداء | سلوك الجدولة الجينية | GeneticScheduler + Demo |
| الأمان | العزل والصلاحيات وXML | Feature + Middleware |

## ج-3 بيئة الاختبار

**جدول ج-2.** بيئة الاختبار

| العنصر | المواصفات |
|--------|-----------|
| نظام التشغيل | Windows 10/11 |
| الخادم | localhost — Laravel |
| قاعدة البيانات | MySQL (اختبار) |
| الواجهة | React SPA (Vite) |
| المتصفحات | Chrome / Edge |

## ج-4 حالات الاختبار (نموذج)

**جدول ج-3.** نموذج حالات الاختبار

| رقم الحالة | الوصف | المدخلات | الخرج المتوقع | النتيجة |
|------------|-------|----------|---------------|---------|
| TC-01 | تسجيل طالب مطابق | إيميل+رقم صحيحان | active | تم |
| TC-02 | تسجيل طالب غير مطابق | بيانات غير موجودة | رفض | تم |
| TC-03 | تسجيل مشرف مطابق | إيميل مطابق | active | تم |
| TC-04 | تقديم مقترح | عنوان+وصف+مشرف | pending | تم |
| TC-05 | اعتماد مقترح | موافقة مشرف | إنشاء مشروع | تم |
| TC-06 | منع تخطي مرحلة | مرحلة لاحقة دون نجاح | رفض/قفل | تم |
| TC-07 | توليد جدول | مشاريع+توافر+قاعات | مرشحون | تم |
| TC-08 | منع تعارض لجنة | مشرف في اللجنة | رفض | تم |
"""


def write_docx(md: str, out: Path) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.shared import Pt, Cm

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2.54)
        sec.bottom_margin = Cm(2.54)
        sec.left_margin = Cm(2.54)
        sec.right_margin = Cm(2.54)

    style = doc.styles["Normal"]
    style.font.name = "Simplified Arabic"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Simplified Arabic")
    style.font.size = Pt(14)
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(6)

    def rtl(p):
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pPr = p._p.get_or_add_pPr()
        for child in list(pPr):
            if child.tag == qn("w:bidi"):
                pPr.remove(child)
        pPr.append(pPr.makeelement(qn("w:bidi"), {}))

    def add(text, bold=False, size=14, center=False):
        p = doc.add_paragraph()
        if center:
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            rtl(p)
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        r.font.name = "Simplified Arabic"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Simplified Arabic")
        return p

    table_buf: list[str] = []
    in_code = False
    code: list[str] = []

    def flush_table():
        nonlocal table_buf
        if not table_buf:
            return
        rows = []
        for row in table_buf:
            if set(row.replace("|", "").strip()) <= set("-: "):
                continue
            rows.append([c.strip() for c in row.strip().strip("|").split("|")])
        table_buf = []
        if not rows:
            return
        cols = max(len(r) for r in rows)
        t = doc.add_table(rows=len(rows), cols=cols)
        t.style = "Table Grid"
        for ri, row in enumerate(rows):
            for ci in range(cols):
                cell = t.cell(ri, ci)
                cell.text = row[ci] if ci < len(row) else ""
                for p in cell.paragraphs:
                    rtl(p)
                    for run in p.runs:
                        run.font.name = "Simplified Arabic"
                        run.font.size = Pt(12)
                        run.bold = ri == 0
        doc.add_paragraph()

    for line in md.splitlines():
        if line.strip().startswith("```"):
            if not in_code:
                in_code, code = True, []
            else:
                in_code = False
                add("\n".join(code), size=11)
            continue
        if in_code:
            code.append(line)
            continue
        if line.strip().startswith("|"):
            table_buf.append(line)
            continue
        flush_table()
        s = line.strip()
        if not s or s == "---":
            continue
        if s.startswith("# "):
            add(s[2:], bold=True, size=18, center=True)
        elif s.startswith("## "):
            add(s[3:], bold=True, size=16)
        elif s.startswith("### "):
            add(s[4:], bold=True, size=14)
        elif s.startswith("- "):
            add("• " + s[2:].replace("**", ""), size=14)
        elif s.startswith("**جدول") or s.startswith("**شكل") or s.startswith("**Table") or s.startswith("**Figure"):
            add(s.replace("**", ""), bold=True, size=14, center=True)
        else:
            add(s.replace("**", "").replace("`", ""), size=14)

    flush_table()
    doc.save(out)
    print("DOCX", out.name)


def pdf_from_html(html_path: Path, pdf_path: Path, page) -> None:
    page.goto(html_path.as_uri())
    page.pdf(
        path=str(pdf_path),
        format="A4",
        print_background=True,
        margin={"top": "0", "bottom": "0", "left": "0", "right": "0"},  # margins in CSS @page
    )


def main() -> None:
    md = md_from_guide()
    (ROOT / "07-chapter-testing-and-evaluation.md").write_text(md, encoding="utf-8")
    print("MD chapter")

    docx_path = ROOT / "07-chapter-testing-and-evaluation.docx"
    try:
        write_docx(md, docx_path)
    except PermissionError:
        alt = ROOT / "07-chapter-testing-and-evaluation-GUIDE.docx"
        write_docx(md, alt)
        print("DOCX locked; wrote", alt.name)

    # Full chapter HTML/PDF
    full_html = wrap("الفصل السابع: الاختبار والتقييم", CHAPTER_BODY)
    full_path = ROOT / "07-chapter-testing-and-evaluation.html"
    full_path.write_text(full_html, encoding="utf-8")

    appendix_html = wrap("ملحق ج — خطة الاختبار", APPENDIX_C)
    appendix_path = ROOT / "appendix-c-test-plan.html"
    appendix_path.write_text(appendix_html, encoding="utf-8")

    # Combined chapter + appendix
    combined = wrap(
        "الفصل السابع والملحق ج",
        CHAPTER_BODY + '<div class="page-break"></div>' + APPENDIX_C,
    )
    combined_path = ROOT / "07-chapter-testing-and-evaluation-with-appendix-c.html"
    combined_path.write_text(combined, encoding="utf-8")

    with sync_playwright() as p:
        browser = p.chromium.launch()
        page = browser.new_page()
        pdf_from_html(full_path, ROOT / "07-chapter-testing-and-evaluation.pdf", page)
        print("PDF chapter")
        pdf_from_html(appendix_path, ROOT / "appendix-c-test-plan.pdf", page)
        print("PDF appendix C")
        pdf_from_html(combined_path, ROOT / "07-chapter-testing-and-evaluation-with-appendix-c.pdf", page)
        print("PDF chapter+appendix")
        browser.close()

    (ROOT / "README.md").write_text(
        """# الفصل السابع — الاختبار والتقييم (وفق دليل د. كادان الجمعة)

التنسيق مطابق للدليل:
- خط عربي: Simplified Arabic — عناوين 18/16/14 غامق، متن 14
- تباعد أسطر 1.5 — هوامش 2.54 سم
- عنوان الجدول **فوق** الجدول: «جدول س. العنوان»
- عنوان الشكل **تحت** الشكل: «شكل س. العنوان»
- أقسام الفصل: 7-1 … 7-6 + **ملحق ج**

## للنسخ

| الملف | الوصف |
|--------|--------|
| `07-chapter-testing-and-evaluation.docx` | Word جاهز (عربي) |
| `07-chapter-testing-and-evaluation.md` | Markdown للنسخ |
| `07-chapter-testing-and-evaluation.pdf` | PDF الفصل |
| `appendix-c-test-plan.pdf` | ملحق ج |
| `07-chapter-testing-and-evaluation-with-appendix-c.pdf` | الفصل + الملحق معاً |

English edition: see `README.en.md` / `*.en.*` files.
""",
        encoding="utf-8",
    )
    print("README")


if __name__ == "__main__":
    main()
