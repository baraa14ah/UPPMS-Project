#!/usr/bin/env python3
"""Chapter 7 EN — formatted per SPU guide (Times New Roman, 2.54cm, 1.5 spacing)."""
from __future__ import annotations

from pathlib import Path

from playwright.sync_api import sync_playwright

ROOT = Path(__file__).resolve().parents[1] / "chapter-07-testing-and-evaluation"
DIAG = ROOT / "diagrams-en"
TOTAL, UNIT, FEATURE = 141, 50, 91

STYLE = """
  @page { size: A4; margin: 2.54cm; }
  body {
    font-family: "Times New Roman", Times, serif;
    font-size: 14pt; line-height: 1.5; color: #000; margin: 0; background: #fff;
    text-align: justify; direction: ltr;
  }
  h1.chapter { font-size: 18pt; font-weight: bold; text-align: center; margin: 0 0 22px; }
  h2 { font-size: 16pt; font-weight: bold; margin: 20px 0 10px; page-break-after: avoid; }
  h3 { font-size: 14pt; font-weight: bold; margin: 14px 0 8px; page-break-after: avoid; }
  p { margin: 0 0 10px; text-indent: 1.2em; }
  p.ni { text-indent: 0; }
  ul, ol { margin: 6px 0 12px; padding-left: 1.5em; }
  li { margin: 4px 0; }
  table {
    width: 100%; border-collapse: collapse; margin: 6px 0 14px;
    font-size: 11pt; page-break-inside: avoid;
  }
  th, td { border: 1px solid #000; padding: 6px 8px; text-align: left; vertical-align: top; }
  th { background: #f0f0f0; font-weight: bold; }
  .tbl-cap { font-size: 14pt; font-weight: bold; text-align: center; margin: 14px 0 6px; text-indent: 0; }
  .fig-cap { font-size: 14pt; font-weight: bold; text-align: center; margin: 6px 0 14px; text-indent: 0; }
  .imgbox { margin: 10px 0 4px; text-align: center; page-break-inside: avoid; }
  .imgbox img { max-width: 95%; height: auto; border: 1px solid #999; }
  .note { margin: 12px 0; padding: 8px 12px; border: 1px solid #555; text-indent: 0; font-size: 12pt; }
  .page-break { page-break-before: always; }
"""


def table(caption: str, headers: list[str], rows: list[list[str]]) -> str:
    head = "".join(f"<th>{h}</th>" for h in headers)
    body = "".join("<tr>" + "".join(f"<td>{c}</td>" for c in row) + "</tr>" for row in rows)
    return f'<p class="tbl-cap">{caption}</p><table><tr>{head}</tr>{body}</table>'


def figure(path: Path, caption: str) -> str:
    if path.exists():
        return f'<div class="imgbox"><img src="{path.as_uri()}" alt="{caption}"/></div><p class="fig-cap">{caption}</p>'
    return f'<p class="note">[Insert figure: {caption}]</p>'


def wrap(title: str, body: str) -> str:
    return f"""<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8"/><title>{title}</title>
<style>{STYLE}</style></head><body>{body}</body></html>"""


CHAPTER = f"""
<h1 class="chapter">Chapter Seven: Testing and Evaluation</h1>
<p>
This chapter aims to prove that the University Project Portfolio Management System (UPPMS)
works as required according to its functional and non-functional requirements.
Accordingly, this chapter presents the test plan, types of tests, test diagrams,
test results and analysis, quality metrics, and user testing outcomes.
Operational test-case details are completed in Appendix C.
</p>

<h2>7-1 Test Plan</h2>
<p>
The test plan answers: What will be tested? When? With which tools? Who is responsible?
(See Appendix C.) Figure 7-1 shows the test-plan flow.
</p>
{figure(DIAG / "01-test-plan-flow.png", "Figure 7-1. Test Plan Flow Diagram")}
<p>Table 7-1 summarizes the test scope for UPPMS.</p>
{table("Table 7-1. Test Scope", ["Item", "Details"], [
    ["What will be tested", "Backend services, DB integration, API routes, university isolation, XML registration, proposals, academic tracks, committees, genetic scheduling, and SPU campus acceptance scenarios"],
    ["What will not be fully tested here", "Exhaustive cross-browser/device compatibility and large-scale production stress testing outside the development environment"],
    ["When", "During each Sprint, at phase end for the full suite, and before delivery for acceptance"],
    ["Tools", "PHPUnit 10 in Laravel, MySQL test DB, and manual verification on the React SPA"],
    ["Responsible", "Development team (automation/analysis), academic supervisor (review), and pilot users (acceptance)"],
])}
{table("Table 7-2. Testing Responsibilities and Schedule", ["Phase", "Owner", "Activity"], [
    ["During each Sprint", "Development team", "Unit and system tests for the related feature"],
    ["End of phase", "Development team", "Full suite + regression"],
    ["Before delivery", "Team + pilot users", "Acceptance and user testing on SPU demo data"],
])}
<p class="ni note">See Appendix C for the detailed test-plan template and expanded cases.</p>

<h2>7-2 Types of Tests</h2>
<p>
As required by the guide, the project applies: unit, integration, system, acceptance,
performance/stress, security, usability, compatibility, and regression testing.
</p>
{table("Table 7-3. Types of Tests Applied in UPPMS", ["Test Type", "Purpose in the Project", "Tool / Approach"], [
    ["Unit", "Fitness, GA operators, Track/XML services", f"PHPUnit Unit ({UNIT} cases)"],
    ["Integration", "Services with database/Eloquent", "Feature/Unit + DatabaseTransactions"],
    ["System", "Full API flows (proposals, tracks, committees, XML, scheduling)", f"PHPUnit Feature ({FEATURE} cases)"],
    ["Acceptance", "Admin / supervisor / student scenarios", "Manual + SpuCampusDemoSeeder"],
    ["Performance / Stress", "Genetic scheduler under denser loads", "GA on Campus Demo"],
    ["Security", "Isolation, RBAC, XML matching, pending users", "Feature tests + Middleware review"],
    ["Usability", "Clarity of key screens", "User testing + notes"],
    ["Compatibility", "Modern browsers and RTL/LTR", "Manual UI checks"],
    ["Regression", "No breakage after fixes", "php vendor/bin/phpunit"],
])}
{figure(DIAG / "02-test-pyramid.png", "Figure 7-2. UPPMS Test Pyramid")}
<p>Total automated cases in the repository: <strong>{TOTAL}</strong> (Unit: {UNIT} — Feature: {FEATURE}).</p>

<h2>7-3 Test Diagrams</h2>
<p>
Test diagrams include Test Case Diagrams and test-results diagrams.
Figure 7-3 shows modules covered by automated tests.
</p>
{figure(DIAG / "03-test-coverage-modules.png", "Figure 7-3. Module Coverage by Automated Tests")}
{table("Table 7-4. Sample Test Cases", ["Case ID", "Description", "Inputs", "Expected Output"], [
    ["TC-XML-01", "Matching student XML registration", "Matching email + university number", "active account"],
    ["TC-XML-02", "Matching supervisor XML registration", "Matching email only", "active account"],
    ["TC-PROP-01", "Submit proposals without early track lock", "Up to 3 pending proposals", "Success without early track assignment"],
    ["TC-PROP-02", "Approve proposal", "Supervisor approval", "Project created + track assigned"],
    ["TC-TRACK-01", "Prevent skipping prerequisites", "Later stage without prior pass", "Reject / keep locked"],
    ["TC-GA-01", "Generate defense schedules", "Projects + availability + rooms", "Valid candidates + fitness"],
    ["TC-COM-01", "Prevent committee conflict of interest", "Project supervisor on own committee", "Conflict rejected"],
])}
{figure(DIAG / "04-results-summary.png", "Figure 7-4. Automated Test Case Distribution")}

<h2>7-4 Test Results and Analysis</h2>
{table("Table 7-5. Automated Test Inventory Summary", ["Metric", "Value"], [
    ["Total test cases", str(TOTAL)], ["Unit", str(UNIT)], ["Feature / System", str(FEATURE)], ["Test files", "20"],
])}
{table("Table 7-6. Coverage by Main Test Files", ["File", "Type", "Focus"], [
    ["ProjectProposalControllerTest", "Feature", "Proposals + isolation"],
    ["TrackControllerTest / TrackServiceTest", "Feature + Unit", "Tracks and progress"],
    ["XmlImport* / XmlRegistrationTest", "Feature + Unit", "XML import/match"],
    ["CommitteeControllerTest", "Feature", "Committees"],
    ["GeneticScheduler* / Fitness* / Evolutionary*", "Unit", "Scheduling engine"],
    ["ScheduleControllerTest / ScheduleApproval*", "Feature + Unit", "Generate/approve"],
    ["AIIdeation* / AITask*", "Unit + Feature", "AI services"],
])}
<ul>
  <li>Passed cases cover critical paths (XML, proposals, tracks, committees, scheduling).</li>
  <li>Failures are fixed immediately, then re-checked with regression.</li>
  <li>No critical remaining defects are accepted on security/isolation before delivery.</li>
  <li>Execution time is recorded from PHPUnit output for formal documentation runs.</li>
</ul>

<h2>7-5 Quality Metrics</h2>
{table("Table 7-7. Quality Metrics Used in UPPMS", ["Metric", "Definition", "Project Measurement"], [
    ["Test Coverage", "Share of critical features covered", "XML, proposals, tracks, committees, GA, AI"],
    ["Defect Density", "Defects ÷ component size", "Logged per Spec Kit phase and re-tested"],
    ["Pass Rate", "Passed ÷ executed", "Delivery target ≥ 95% on critical paths"],
    ["Response Time", "API latency for common operations", "CRUD interactive; GA longer with input size"],
    ["Error Rate", "Incorrect responses / unhandled exceptions", "Via failing tests and Laravel logs"],
    ["Security Isolation", "No cross-university leakage", "TenantScope + multi-university Feature tests"],
])}
<p>
Acceptance criteria: no security/isolation/XML failures; successful Admin/Supervisor/Student acceptance;
ability to generate and approve a defense schedule on Campus Demo without visible hard conflicts.
</p>

<h2>7-6 User Testing</h2>
<p>
User testing used Syrian Private University demo data via <code>SpuCampusDemoSeeder</code>
(28 students, 10 supervisors, 28 projects, 4 committees).
</p>
{table("Table 7-8. User Testing Sample and Scenarios", ["Role", "Demo Account", "Scenario"], [
    ["University Admin", "spu-campus-admin@syrian-private.local", "XML, tracks, committees, schedule generate/approve"],
    ["Supervisor", "spu-campus-supervisor-01@syrian-private.local", "Proposals, availability, projects"],
    ["Student", "spu-campus-student-01@syrian-private.local", "Proposal, progress, tasks, notifications"],
])}
<p class="ni">Shared demo password: <code>password</code>.</p>
{table("Table 7-9. User Feedback and Improvements", ["Observation", "Impact", "Improvement"], [
    ["Early track lock on proposal submit", "Stage-status confusion", "Defer assignment until approval (done)"],
    ["Multi-university supervisors missing membership proposals", "Delayed review", "Extend TenantScope (done)"],
    ["Need for ready demo data", "Slower demos", "SpuCampusDemoSeeder"],
])}
<p>
In conclusion, primary flows are usable by all three roles, and critical isolation/track-lock issues
were addressed before final delivery.
</p>
"""

APPENDIX = f"""
<h1 class="chapter">Appendix C — Test Plan Template (UPPMS)</h1>
<p class="ni">This appendix follows the guide’s Appendix C model, adapted to UPPMS.</p>
<h2>C-1 Test Scope</h2>
<ul>
  <li><strong>In scope:</strong> backend units, integration, API, isolation, scheduling performance, security.</li>
  <li><strong>Out of full scope:</strong> exhaustive manual UI and all-browser compatibility.</li>
</ul>
<h2>C-2 Test Strategy</h2>
{table("Table C-1. Test Strategy", ["Test Type", "Objective", "Tool"], [
    ["Unit", "Each function/unit separately", "PHPUnit Unit"],
    ["Integration", "Units with DB", "PHPUnit + Transactions"],
    ["System", "Full API flows", "PHPUnit Feature"],
    ["Performance", "GA under denser data", "GeneticScheduler + Demo"],
    ["Security", "Isolation, RBAC, XML", "Feature + Middleware"],
])}
<h2>C-3 Test Environment</h2>
{table("Table C-2. Test Environment", ["Item", "Specification"], [
    ["OS", "Windows 10/11"], ["Server", "localhost — Laravel"],
    ["Database", "MySQL (test)"], ["UI", "React SPA (Vite)"], ["Browsers", "Chrome / Edge"],
])}
<h2>C-4 Test Cases (Sample)</h2>
{table("Table C-3. Sample Test Cases", ["ID", "Description", "Inputs", "Expected", "Result"], [
    ["TC-01", "Matching student registration", "Valid XML email+number", "active", "Done"],
    ["TC-02", "Non-matching student registration", "Unknown credentials", "Reject", "Done"],
    ["TC-03", "Matching supervisor registration", "Matching email", "active", "Done"],
    ["TC-04", "Submit proposal", "Title+description+supervisor", "pending", "Done"],
    ["TC-05", "Approve proposal", "Supervisor approve", "Project created", "Done"],
    ["TC-06", "Block stage skip", "Later stage without prior pass", "Reject/lock", "Done"],
    ["TC-07", "Generate schedule", "Projects+availability+rooms", "Candidates", "Done"],
    ["TC-08", "Committee conflict", "Supervisor on own committee", "Reject", "Done"],
])}
"""

MD = f"""# Chapter Seven: Testing and Evaluation

This chapter aims to prove that UPPMS works as required. It presents the test plan, types of tests, test diagrams, results and analysis, quality metrics, and user testing. Details are completed in Appendix C.

## 7-1 Test Plan

What / When / Tools / Who? (See Appendix C).

**Figure 7-1.** Test Plan Flow Diagram — `diagrams-en/01-test-plan-flow.png`

**Table 7-1.** Test Scope

| Item | Details |
|------|---------|
| What will be tested | Backend services, DB integration, API, isolation, XML, proposals, tracks, committees, genetic scheduling, SPU acceptance |
| What will not be fully tested here | Exhaustive multi-browser compatibility and large-scale production stress |
| When | Each Sprint, phase end, before delivery |
| Tools | PHPUnit 10, MySQL, React manual checks |
| Responsible | Dev team, supervisor, pilot users |

**Table 7-2.** Testing Responsibilities and Schedule

| Phase | Owner | Activity |
|-------|-------|----------|
| Sprint | Development team | Feature-related unit/system tests |
| Phase end | Development team | Full suite + regression |
| Before delivery | Team + pilots | Acceptance / user testing |

## 7-2 Types of Tests

Unit, Integration, System, Acceptance, Performance/Stress, Security, Usability, Compatibility, Regression.

**Table 7-3.** Types of Tests Applied in UPPMS

| Test Type | Purpose | Tool / Approach |
|-----------|---------|-----------------|
| Unit | Fitness/GA/Track/XML units | PHPUnit Unit ({UNIT}) |
| Integration | Services + DB | Feature/Unit + Transactions |
| System | Full API flows | PHPUnit Feature ({FEATURE}) |
| Acceptance | Admin/Supervisor/Student | Manual + SpuCampusDemoSeeder |
| Performance/Stress | GA under denser load | Campus Demo GA runs |
| Security | Isolation/RBAC/XML | Feature + Middleware |
| Usability | UI clarity | User testing |
| Compatibility | Chrome/Edge, RTL/LTR | Manual |
| Regression | No breakage | phpunit |

**Figure 7-2.** UPPMS Test Pyramid — `diagrams-en/02-test-pyramid.png`

Total automated cases: **{TOTAL}** (Unit: {UNIT} — Feature: {FEATURE}).

## 7-3 Test Diagrams

**Figure 7-3.** Module Coverage — `diagrams-en/03-test-coverage-modules.png`

**Table 7-4.** Sample Test Cases

| Case ID | Description | Inputs | Expected Output |
|---------|-------------|--------|-----------------|
| TC-XML-01 | Matching student XML registration | email + number | active |
| TC-XML-02 | Matching supervisor XML registration | email only | active |
| TC-PROP-01 | Proposals without early track lock | up to 3 pending | no early track |
| TC-PROP-02 | Approve proposal | supervisor approve | project + track |
| TC-TRACK-01 | Prevent prerequisite skip | later stage | reject/lock |
| TC-GA-01 | Generate schedules | projects+availability+rooms | candidates + fitness |
| TC-COM-01 | Committee conflict | supervisor on own committee | reject |

**Figure 7-4.** Automated Test Distribution — `diagrams-en/04-results-summary.png`

## 7-4 Test Results and Analysis

**Table 7-5.** Inventory — Total {TOTAL}, Unit {UNIT}, Feature {FEATURE}, 20 files.

**Table 7-6.** Main files: ProjectProposal*, Track*, Xml*, Committee*, Genetic/Fitness/Evolutionary*, Schedule*, AI*.

Analysis: critical paths covered; failures fixed then regression; no critical isolation/security leftovers before delivery; execution time recorded from PHPUnit.

## 7-5 Quality Metrics

**Table 7-7.** Test Coverage, Defect Density, Pass Rate (≥95% target), Response Time, Error Rate, Security Isolation.

## 7-6 User Testing

SPU demo (`SpuCampusDemoSeeder`): 28 students, 10 supervisors, 28 projects, 4 committees. Password: `password`.

**Table 7-8.** Admin / Supervisor / Student demo accounts and scenarios.

**Table 7-9.** Feedback: early track lock (fixed), multi-uni supervisor visibility (fixed), demo data seeder.

---

# Appendix C — Test Plan Template (UPPMS)

## C-1 Scope / C-2 Strategy / C-3 Environment / C-4 Sample Cases
(See PDF/DOCX for full tables TC-01…TC-08.)
"""


def write_docx(md: str, out: Path) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.shared import Pt, Cm

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Cm(2.54)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(14)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    def add(text, bold=False, size=14, center=False):
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        r.font.name = "Times New Roman"
        return p

    table_buf = []

    def flush():
        nonlocal table_buf
        if not table_buf:
            return
        rows = []
        for row in table_buf:
            if set(row.replace("|", "").strip()) <= set("-: "):
                continue
            rows.append([c.strip() for c in row.strip().strip("|").split("|")])
        table_buf = []
        if not rows:
            return
        cols = max(len(r) for r in rows)
        t = doc.add_table(rows=len(rows), cols=cols)
        t.style = "Table Grid"
        for ri, row in enumerate(rows):
            for ci in range(cols):
                cell = t.cell(ri, ci)
                cell.text = row[ci] if ci < len(row) else ""
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(11)
                        run.bold = ri == 0
        doc.add_paragraph()

    for line in md.splitlines():
        if line.strip().startswith("|"):
            table_buf.append(line)
            continue
        flush()
        s = line.strip()
        if not s or s == "---":
            continue
        if s.startswith("# "):
            add(s[2:], True, 18, True)
        elif s.startswith("## "):
            add(s[3:], True, 16)
        elif s.startswith("### "):
            add(s[4:], True, 14)
        elif s.startswith("- "):
            add("• " + s[2:].replace("**", ""))
        elif s.startswith("**Table") or s.startswith("**Figure") or s.startswith("**جدول") or s.startswith("**شكل"):
            add(s.replace("**", ""), True, 14, True)
        else:
            add(s.replace("**", "").replace("`", ""))
    flush()
    doc.save(out)
    print("DOCX", out.name)


def main() -> None:
    # Ensure EN diagrams exist
    if not (DIAG / "01-test-plan-flow.png").exists():
        import subprocess, sys
        subprocess.check_call([sys.executable, str(Path(__file__).parent / "generate_chapter07_diagrams_en.py")])

    md = MD
    (ROOT / "07-chapter-testing-and-evaluation.en.md").write_text(md, encoding="utf-8")
    docx_path = ROOT / "07-chapter-testing-and-evaluation.en.docx"
    try:
        write_docx(md, docx_path)
    except PermissionError:
        alt = ROOT / "07-chapter-testing-and-evaluation-GUIDE.en.docx"
        write_docx(md, alt)
        print("DOCX locked; wrote", alt.name)

    full = wrap("Chapter Seven", CHAPTER)
    app = wrap("Appendix C", APPENDIX)
    both = wrap("Chapter Seven + Appendix C", CHAPTER + '<div class="page-break"></div>' + APPENDIX)
    fp = ROOT / "07-chapter-testing-and-evaluation.en.html"
    ap = ROOT / "appendix-c-test-plan.en.html"
    bp = ROOT / "07-chapter-testing-and-evaluation-with-appendix-c.en.html"
    fp.write_text(full, encoding="utf-8")
    ap.write_text(app, encoding="utf-8")
    bp.write_text(both, encoding="utf-8")

    with sync_playwright() as p:
        browser = p.chromium.launch()
        page = browser.new_page()
        for html, pdf in [
            (fp, ROOT / "07-chapter-testing-and-evaluation.en.pdf"),
            (ap, ROOT / "appendix-c-test-plan.en.pdf"),
            (bp, ROOT / "07-chapter-testing-and-evaluation-with-appendix-c.en.pdf"),
        ]:
            page.goto(html.as_uri())
            page.pdf(path=str(pdf), format="A4", print_background=True,
                     margin={"top": "0", "bottom": "0", "left": "0", "right": "0"})
            print("PDF", pdf.name)
        browser.close()

    (ROOT / "README.en.md").write_text(
        """# Chapter Seven — Testing and Evaluation (English, per SPU Guide)

Formatting per Dr. Kadan Al-Jumaa guide:
- Times New Roman — 18/16/14 bold headings, 14 body
- Line spacing 1.5 — margins 2.54 cm
- Table caption **above**: “Table X. Title”
- Figure caption **below**: “Figure X. Title”
- Sections 7-1 … 7-6 + **Appendix C**

## Copy files

| File | Description |
|------|-------------|
| `07-chapter-testing-and-evaluation.en.docx` | Word |
| `07-chapter-testing-and-evaluation.en.md` | Markdown |
| `07-chapter-testing-and-evaluation.en.pdf` | Chapter PDF |
| `appendix-c-test-plan.en.pdf` | Appendix C |
| `07-chapter-testing-and-evaluation-with-appendix-c.en.pdf` | Chapter + Appendix |
""",
        encoding="utf-8",
    )
    print("README.en.md")


if __name__ == "__main__":
    main()
