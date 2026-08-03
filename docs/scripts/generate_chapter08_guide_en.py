#!/usr/bin/env python3
"""
Chapter 8 — Results and Analysis (English only)
Per SPU guide (Dr. Kadan Al-Jumaa):
  8-1 System application results (screenshots, reports, stats, charts)
  8-2 Comparative analysis with existing systems
Leave blanks for screenshots / unknown data.
Formatting: Times New Roman, 18/16/14, 1.5 spacing, margins 2.54cm
"""
from __future__ import annotations

from pathlib import Path

from playwright.sync_api import sync_playwright

ROOT = Path(__file__).resolve().parents[1] / "chapter-08-results-and-analysis"
EN = ROOT / "en"
COPY = EN / "copy"
DIAG = EN / "diagrams"
PLACE = EN / "placeholders"
for d in (COPY, DIAG, PLACE):
    d.mkdir(parents=True, exist_ok=True)

STYLE = """
  @page { size: A4; margin: 2.54cm; }
  body {
    font-family: "Times New Roman", Times, serif;
    font-size: 14pt; line-height: 1.5; color: #000; margin: 0; background: #fff;
    text-align: justify; direction: ltr;
  }
  h1.chapter { font-size: 18pt; font-weight: bold; text-align: center; margin: 0 0 22px; }
  h2 { font-size: 16pt; font-weight: bold; margin: 20px 0 10px; page-break-after: avoid; }
  h3 { font-size: 14pt; font-weight: bold; margin: 14px 0 8px; page-break-after: avoid; }
  p { margin: 0 0 10px; text-indent: 1.2em; }
  p.ni { text-indent: 0; }
  ul, ol { margin: 6px 0 12px; padding-left: 1.5em; }
  li { margin: 4px 0; }
  table {
    width: 100%; border-collapse: collapse; margin: 6px 0 14px;
    font-size: 11pt; page-break-inside: avoid;
  }
  th, td { border: 1px solid #000; padding: 6px 8px; text-align: left; vertical-align: top; }
  th { background: #f0f0f0; font-weight: bold; }
  .tbl-cap { font-size: 14pt; font-weight: bold; text-align: center; margin: 14px 0 6px; text-indent: 0; }
  .fig-cap { font-size: 14pt; font-weight: bold; text-align: center; margin: 6px 0 14px; text-indent: 0; }
  .ph {
    margin: 10px auto 4px; width: 95%; min-height: 160px;
    border: 2px dashed #888; background: #fafafa;
    display: flex; align-items: center; justify-content: center;
    text-align: center; color: #555; font-size: 12pt; page-break-inside: avoid;
    padding: 16px;
  }
  .blank { color: #666; font-style: italic; }
  .note { margin: 12px 0; padding: 8px 12px; border: 1px solid #555; text-indent: 0; font-size: 12pt; }
"""


def table(caption: str, headers: list[str], rows: list[list[str]]) -> str:
    head = "".join(f"<th>{h}</th>" for h in headers)
    body = "".join("<tr>" + "".join(f"<td>{c}</td>" for c in row) + "</tr>" for row in rows)
    return f'<p class="tbl-cap">{caption}</p><table><tr>{head}</tr>{body}</table>'


def screenshot_slot(fig_no: str, title: str, hint: str) -> str:
    return f"""
<div class="ph">
  [SCREENSHOT PLACEHOLDER]<br/><br/>
  Insert screenshot here:<br/><strong>{title}</strong><br/><br/>
  <span style="font-size:11pt">{hint}</span>
</div>
<p class="fig-cap">Figure {fig_no}. {title}</p>
"""


def blank_cell() -> str:
    return '<span class="blank">[ ]</span>'


CHAPTER = f"""
<h1 class="chapter">Chapter Eight: Results and Analysis</h1>

<p>
This chapter aims to present what the system actually produced in practice.
Accordingly, it includes the results of applying UPPMS (screenshots, reports,
statistics, and charts that illustrate system performance), followed by a
comparative analysis with existing systems in terms of speed, accuracy, cost,
ease of use, and features.
</p>

<p class="ni note">
Blank boxes marked <strong>[SCREENSHOT PLACEHOLDER]</strong> and cells marked
<span class="blank">[ ]</span> are left empty intentionally for the author to
fill with real captures, measured values, or external references when available.
</p>

<h2>8-1 System Application Results</h2>

<p>
This section presents the outcomes of running UPPMS on the Syrian Private University
demo environment. The results are organized as screenshots of the main flows,
summary reports/statistics, and charts of system performance.
</p>

<h3>8-1-1 Screenshots of Main System Flows</h3>

<p>
The following figures illustrate the principal end-to-end flows of UPPMS.
Each figure should be replaced with an actual screen capture from the running system.
</p>

{screenshot_slot("8-1", "User Registration / Login",
                 "Show student/supervisor registration or login screen (XML-matched account preferred).")}

{screenshot_slot("8-2", "University Admin Dashboard",
                 "Show admin home / university workspace after login.")}

{screenshot_slot("8-3", "XML Authorized-Users Import",
                 "Show Users → XML Import with an uploaded file and import result summary.")}

{screenshot_slot("8-4", "Project Proposal Submission (Student)",
                 "Show proposal form and list of submitted proposals.")}

{screenshot_slot("8-5", "Supervisor Proposal Review / Approval",
                 "Show supervisor view of a pending proposal with approve/reject actions.")}

{screenshot_slot("8-6", "Academic Track / Student Progress Timeline",
                 "Show track stages (locked / in progress / passed) for a demo student.")}

{screenshot_slot("8-7", "AI Project Ideation",
                 "Show AI idea suggestions screen (Gemini) if used in the demo.")}

{screenshot_slot("8-8", "Project Workspace / Tasks",
                 "Show an approved project with tasks (optionally AI-generated).")}

{screenshot_slot("8-9", "Defense Committee Management",
                 "Show committee list/members for SPU demo committees.")}

{screenshot_slot("8-10", "Genetic Scheduling Dashboard",
                 "Show schedule generation candidates and/or an approved defense schedule.")}

<h3>8-1-2 Reports and Statistics</h3>

<p>
Table 8-1 summarizes quantitative results observed from the demo deployment.
Values left blank should be filled after a measured run (seeded demo counts may
be used where already known).
</p>

{table(
    "Table 8-1. System Application Statistics (Demo / Measured)",
    ["Metric", "Value", "Notes"],
    [
        ["University under test", "syrian private uni (SPU)", "Demo tenant"],
        ["Active students (demo)", "28", "SpuCampusDemoSeeder"],
        ["Active supervisors (demo)", "10", "SpuCampusDemoSeeder"],
        ["Active projects (demo)", "28", "Approved proposals"],
        ["Defense committees (demo)", "4", "Seminar / technical / final"],
        ["Available rooms (demo)", "7", "Including premium rooms"],
        ["Automated test cases in repository", "141", "50 Unit + 91 Feature"],
        ["Average API response time (common CRUD)", blank_cell(), "Measure and fill"],
        ["Schedule generation time (GA, demo load)", blank_cell(), "Measure and fill"],
        ["Proposal approval turnaround (demo observation)", blank_cell(), "Optional"],
        ["User-satisfaction score (survey)", blank_cell(), "If a survey was conducted"],
    ],
)}

<p>
Table 8-2 is reserved for any exported/system-generated reports (PDF exports,
schedule reports, progress summaries). Attach the report file or paste a
screenshot of the report view.
</p>

{table(
    "Table 8-2. Generated Reports Produced by the System",
    ["Report", "Produced?", "Where shown / file name"],
    [
        ["Approved defense schedule view", blank_cell(), blank_cell()],
        ["Student progress / track summary", blank_cell(), blank_cell()],
        ["XML import comparison summary", blank_cell(), blank_cell()],
        ["Other: " + blank_cell(), blank_cell(), blank_cell()],
    ],
)}

<h3>8-1-3 Charts Illustrating System Performance</h3>

<p>
The following chart placeholders should be replaced with actual charts
(e.g., Excel, Python, or dashboard exports) once measurements are available.
</p>

{screenshot_slot("8-11", "Performance Chart — Response Time / Load",
                 "Replace with a chart of API latency or scheduling runtime vs. input size.")}

{screenshot_slot("8-12", "Usage / Outcome Chart",
                 "Replace with a chart of proposals, approvals, defenses, or progress statuses.")}

<p>
A short textual analysis of the charts should be written here after the figures
are inserted:
</p>
<p class="blank ni">[Analysis of Figure 8-11 and Figure 8-12 — to be completed by the author.]</p>

<h2>8-2 Comparative Analysis with Existing Systems</h2>

<p>
This section compares UPPMS with existing approaches/systems used for managing
graduation projects (manual tools, spreadsheets, generic PMS, or institutional
systems). Criteria follow the guide: speed, accuracy, cost, ease of use, and features.
</p>

<p>
Fill the competitor columns with the real systems or approaches you compared against
(names + references). Leave cells blank if a fair measurement is not available.
</p>

{table(
    "Table 8-3. Comparative Analysis: UPPMS vs Existing Systems",
    ["Criterion", "UPPMS", "Existing System / Approach A<br/><span class='blank'>[Name / Ref]</span>", "Existing System / Approach B<br/><span class='blank'>[Name / Ref]</span>"],
    [
        ["Speed (scheduling / coordination)", "Automated genetic scheduling + admin approval", blank_cell(), blank_cell()],
        ["Accuracy (conflicts, prerequisites, isolation)", "Hard/soft constraints, track prerequisites, tenant isolation, XML matching", blank_cell(), blank_cell()],
        ["Cost", blank_cell() + " (hosting / licenses — fill)", blank_cell(), blank_cell()],
        ["Ease of use", "Role-based SPA (Admin / Supervisor / Student), bilingual UI", blank_cell(), blank_cell()],
        ["Core features", "Multi-tenant PMS, XML registration, proposals, AI ideation/tasks, tracks, committees, GA scheduling", blank_cell(), blank_cell()],
        ["Multi-university isolation", "Yes (TenantScope)", blank_cell(), blank_cell()],
        ["Academic track enforcement", "Yes", blank_cell(), blank_cell()],
        ["AI assistance", "Yes (Gemini ideation + task breakdown)", blank_cell(), blank_cell()],
    ],
)}

<p>
Discussion of the comparison (to be completed after filling Table 8-3):
</p>
<p class="blank ni">[Comparative discussion — strengths of UPPMS, limitations, and where existing systems remain preferable.]</p>

<p>
References used for the comparison (if any):
</p>
<p class="blank ni">[Reference 1]</p>
<p class="blank ni">[Reference 2]</p>
<p class="blank ni">[Reference 3]</p>
"""

MD = """# Chapter Eight: Results and Analysis

This chapter aims to present what the system actually produced in practice. It includes system application results (screenshots, reports, statistics, and charts) and a comparative analysis with existing systems (speed, accuracy, cost, ease of use, features).

> Blank markers: `[SCREENSHOT PLACEHOLDER]` and `[ ]` are intentional — fill with real captures, measurements, or references.

## 8-1 System Application Results

### 8-1-1 Screenshots of Main System Flows

**Figure 8-1.** User Registration / Login — `[SCREENSHOT PLACEHOLDER]`

**Figure 8-2.** University Admin Dashboard — `[SCREENSHOT PLACEHOLDER]`

**Figure 8-3.** XML Authorized-Users Import — `[SCREENSHOT PLACEHOLDER]`

**Figure 8-4.** Project Proposal Submission (Student) — `[SCREENSHOT PLACEHOLDER]`

**Figure 8-5.** Supervisor Proposal Review / Approval — `[SCREENSHOT PLACEHOLDER]`

**Figure 8-6.** Academic Track / Student Progress Timeline — `[SCREENSHOT PLACEHOLDER]`

**Figure 8-7.** AI Project Ideation — `[SCREENSHOT PLACEHOLDER]`

**Figure 8-8.** Project Workspace / Tasks — `[SCREENSHOT PLACEHOLDER]`

**Figure 8-9.** Defense Committee Management — `[SCREENSHOT PLACEHOLDER]`

**Figure 8-10.** Genetic Scheduling Dashboard — `[SCREENSHOT PLACEHOLDER]`

### 8-1-2 Reports and Statistics

**Table 8-1.** System Application Statistics (Demo / Measured)

| Metric | Value | Notes |
|--------|-------|-------|
| University under test | syrian private uni (SPU) | Demo tenant |
| Active students (demo) | 28 | SpuCampusDemoSeeder |
| Active supervisors (demo) | 10 | SpuCampusDemoSeeder |
| Active projects (demo) | 28 | Approved proposals |
| Defense committees (demo) | 4 | Seminar / technical / final |
| Available rooms (demo) | 7 | Including premium rooms |
| Automated test cases in repository | 141 | 50 Unit + 91 Feature |
| Average API response time (common CRUD) | [ ] | Measure and fill |
| Schedule generation time (GA, demo load) | [ ] | Measure and fill |
| Proposal approval turnaround | [ ] | Optional |
| User-satisfaction score (survey) | [ ] | If conducted |

**Table 8-2.** Generated Reports Produced by the System

| Report | Produced? | Where shown / file name |
|--------|-----------|-------------------------|
| Approved defense schedule view | [ ] | [ ] |
| Student progress / track summary | [ ] | [ ] |
| XML import comparison summary | [ ] | [ ] |
| Other | [ ] | [ ] |

### 8-1-3 Charts Illustrating System Performance

**Figure 8-11.** Performance Chart — Response Time / Load — `[SCREENSHOT PLACEHOLDER]`

**Figure 8-12.** Usage / Outcome Chart — `[SCREENSHOT PLACEHOLDER]`

[Analysis of Figure 8-11 and Figure 8-12 — to be completed by the author.]

## 8-2 Comparative Analysis with Existing Systems

Criteria: speed, accuracy, cost, ease of use, and features (per the guide).

**Table 8-3.** Comparative Analysis: UPPMS vs Existing Systems

| Criterion | UPPMS | Existing System / Approach A [Name / Ref] | Existing System / Approach B [Name / Ref] |
|-----------|-------|-------------------------------------------|-------------------------------------------|
| Speed (scheduling / coordination) | Automated genetic scheduling + admin approval | [ ] | [ ] |
| Accuracy (conflicts, prerequisites, isolation) | Hard/soft constraints, track prerequisites, tenant isolation, XML matching | [ ] | [ ] |
| Cost | [ ] (hosting / licenses — fill) | [ ] | [ ] |
| Ease of use | Role-based SPA (Admin / Supervisor / Student), bilingual UI | [ ] | [ ] |
| Core features | Multi-tenant PMS, XML registration, proposals, AI ideation/tasks, tracks, committees, GA scheduling | [ ] | [ ] |
| Multi-university isolation | Yes (TenantScope) | [ ] | [ ] |
| Academic track enforcement | Yes | [ ] | [ ] |
| AI assistance | Yes (Gemini ideation + task breakdown) | [ ] | [ ] |

[Comparative discussion — to be completed after filling Table 8-3.]

References:

- [Reference 1]
- [Reference 2]
- [Reference 3]
"""


def write_docx(md: str, out: Path) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.oxml import OxmlElement

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Cm(2.54)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(14)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    def add(text, bold=False, size=14, center=False, italic=False, color=None):
        p = doc.add_paragraph()
        p.paragraph_format.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        )
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        r.font.name = "Times New Roman"
        if color:
            r.font.color.rgb = color
        return p

    def add_placeholder_box(title: str, fig: str):
        # dashed-like note as paragraph block
        p = add(f"[SCREENSHOT PLACEHOLDER]  —  Insert: {title}", True, 12, True, False, RGBColor(0x66, 0x66, 0x66))
        add(f"Figure {fig}. {title}", True, 14, True)

    table_buf: list[str] = []

    def flush():
        nonlocal table_buf
        if not table_buf:
            return
        rows = []
        for row in table_buf:
            if set(row.replace("|", "").strip()) <= set("-: "):
                continue
            rows.append([c.strip() for c in row.strip().strip("|").split("|")])
        table_buf = []
        if not rows:
            return
        cols = max(len(r) for r in rows)
        t = doc.add_table(rows=len(rows), cols=cols)
        t.style = "Table Grid"
        for ri, row in enumerate(rows):
            for ci in range(cols):
                cell = t.cell(ri, ci)
                cell.text = row[ci] if ci < len(row) else ""
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(10)
                        run.bold = ri == 0
        doc.add_paragraph()

    for line in md.splitlines():
        if line.strip().startswith("|"):
            table_buf.append(line)
            continue
        flush()
        s = line.strip()
        if not s or s == "---":
            continue
        if s.startswith("> "):
            add(s[2:].replace("**", ""), False, 12, False, True)
            continue
        if s.startswith("# "):
            add(s[2:], True, 18, True)
        elif s.startswith("## "):
            add(s[3:], True, 16)
        elif s.startswith("### "):
            add(s[4:], True, 14)
        elif s.startswith("- "):
            add("• " + s[2:].replace("**", "").replace("`", ""))
        elif s.startswith("**Figure") and "SCREENSHOT" in s:
            # **Figure 8-1.** Title — `[SCREENSHOT PLACEHOLDER]`
            title = s.replace("**", "").split("—")[0].strip()
            # extract fig number and title
            # Figure 8-1. User Registration / Login
            inner = title.replace("Figure ", "")
            fig_no = inner.split(".", 1)[0].strip()
            fig_title = inner.split(".", 1)[1].strip() if "." in inner else inner
            add_placeholder_box(fig_title, fig_no)
        elif s.startswith("**Table") or s.startswith("**Figure"):
            add(s.replace("**", "").replace("`", ""), True, 14, True)
        else:
            add(s.replace("**", "").replace("`", ""), italic=("[" in s and "]" in s and "to be completed" in s.lower() or s.startswith("[")))

    flush()
    try:
        doc.save(out)
    except PermissionError:
        alt = out.with_name(out.stem + "-NEW.docx")
        doc.save(alt)
        print("DOCX locked; wrote", alt.name)
        return
    print("DOCX", out.name)


def main() -> None:
    md_path = COPY / "chapter-08.md"
    md_path.write_text(MD, encoding="utf-8")
    print("MD", md_path)

    write_docx(MD, COPY / "chapter-08.docx")

    # placeholder note files in diagrams folder
    (DIAG / "README.md").write_text(
        """# Diagrams / Screenshots for Chapter 8

Replace each placeholder below with a real PNG screenshot, then update the chapter figures.

Suggested filenames:

| Figure | Suggested file | What to capture |
|--------|----------------|-----------------|
| 8-1 | `01-login-register.png` | Registration / login |
| 8-2 | `02-admin-dashboard.png` | Admin dashboard |
| 8-3 | `03-xml-import.png` | XML import |
| 8-4 | `04-proposal-submit.png` | Student proposal |
| 8-5 | `05-proposal-review.png` | Supervisor review |
| 8-6 | `06-track-progress.png` | Academic track timeline |
| 8-7 | `07-ai-ideation.png` | AI ideation |
| 8-8 | `08-project-tasks.png` | Project / tasks |
| 8-9 | `09-committees.png` | Committees |
| 8-10 | `10-scheduling.png` | Scheduling dashboard |
| 8-11 | `11-performance-chart.png` | Performance chart |
| 8-12 | `12-usage-chart.png` | Usage / outcomes chart |
""",
        encoding="utf-8",
    )

    html = f"""<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8"/>
<title>Chapter Eight: Results and Analysis</title><style>{STYLE}</style></head>
<body>{CHAPTER}</body></html>"""
    html_path = COPY / "chapter-08.html"
    html_path.write_text(html, encoding="utf-8")

    with sync_playwright() as p:
        browser = p.chromium.launch()
        page = browser.new_page()
        page.goto(html_path.as_uri())
        pdf_path = COPY / "chapter-08.pdf"
        page.pdf(
            path=str(pdf_path),
            format="A4",
            print_background=True,
            margin={"top": "0", "bottom": "0", "left": "0", "right": "0"},
        )
        print("PDF", pdf_path.name)
        browser.close()

    (EN / "README.md").write_text(
        """# Chapter 8 — Results and Analysis (English)

Per SPU guide (Dr. Kadan Al-Jumaa): show what the system actually produced.

## Structure

```
en/
├── copy/          ← start here (Word / Markdown / PDF)
├── diagrams/      ← put real screenshots here
└── README.md
```

## Sections

| Section | Content |
|---------|---------|
| 8-1 System Application Results | Screenshots, reports/statistics, performance charts |
| 8-2 Comparative Analysis | Comparison table: speed, accuracy, cost, ease of use, features |

## Copy files

| File | Use |
|------|-----|
| `copy/chapter-08.docx` | Word (fill blanks + paste screenshots) |
| `copy/chapter-08.md` | Markdown |
| `copy/chapter-08.pdf` | PDF with dashed screenshot placeholders |

## What is left blank on purpose

- All screenshot figures (8-1 … 8-12)
- Measured timings / survey scores
- Competitor system names, values, and references
- Comparative discussion text
""",
        encoding="utf-8",
    )

    (ROOT / "README.md").write_text(
        """# Chapter Eight — Results and Analysis

English-only package (per SPU graduation report guide).

**Start here:** [`en/copy/`](en/copy/)

- `chapter-08.docx` — best for filling screenshots and blanks
- `chapter-08.pdf` — preview with placeholders
- `chapter-08.md` — plain text copy

Screenshots go into [`en/diagrams/`](en/diagrams/) (see README there for filenames).
""",
        encoding="utf-8",
    )
    print("README done")


if __name__ == "__main__":
    main()
